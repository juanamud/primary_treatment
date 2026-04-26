"""
Diseño de PTAR — Caso de estudio: Barbosa, Antioquia
Materia: Sistemas de Tratamiento de Aguas Residuales — UdeA

Tren de tratamiento (3 alternativas para el primario):
  Primario  →  Facultativa  →  Maduración
  ├── Sedimentador 1° (Qasim, RAS 2017)
  ├── Laguna Anaerobia (Mara 1997, von Sperling)
  └── Reactor UASB (van Haandel & Lettinga 1994)

Funcionalidades:
  • Selección de primario por área disponible
  • Planos esquemáticos (planta + corte) por unidad
  • Diagrama del tren completo
  • Manejo de lodos (volumen, evacuación, lechos de secado)
  • Cumplimiento Resolución 0631 de 2015 (Colombia)
  • Análisis de sensibilidad (T, DBO_in)
  • Costos comparativos (US$/hab) según Clase 09
  • Verificación con QMH para UASB
  • Unidades en paralelo
  • Aprovechamiento energético del biogás (kWh)
  • Exportación de memoria a Word (.docx)
"""

import io
import math

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon, Rectangle, FancyBboxPatch

import numpy as np
import pandas as pd
import streamlit as st

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


# ============================================================
# 1. CONFIGURACIÓN
# ============================================================
st.set_page_config(
    page_title="PTAR Barbosa - Diseño completo",
    layout="wide",
    page_icon="💧",
)


# ============================================================
# 2. DATOS ESTÁTICOS DEL CASO DE ESTUDIO
# ============================================================
DATOS_BARBOSA = {
    "municipio": "Barbosa, Antioquia",
    "anio_actual": 2025,
    "horizonte_diseno": 25,
    "poblacion_actual": 55_000,
    "tasa_crecimiento": 1.4,
    "dotacion_neta": 130,
    "factor_retorno": 0.85,
    "k1": 1.30,
    "k2": 1.45,
    "DBO5": 280.0,
    "DQO": 450.0,
    "SST": 220.0,
    "coliformes_fecales": 1.5e7,
    "T_aire_mes_frio": 19.0,
    "area_disponible": 60_000.0,
}

# Resolución 0631 de 2015 — PTAR doméstica > 625 hab equivalentes
RES_0631 = {
    "DBO5_max": 90.0,
    "DQO_max": 180.0,
    "SST_max": 90.0,
    "pH_min": 6.0,
    "pH_max": 9.0,
    "grasas_max": 20.0,
    "S_sed_max": 5.0,
}

# Decreto 1076 de 2015 — Reúso de aguas residuales
REUSO_LIMITES = {
    "Riego restringido": 1000,        # NMP/100mL coliformes
    "Riego no restringido": 100,
    "Vertimiento a cuerpo receptor": 1e6,
}

# Costos referenciales (US$/hab) — Clase 09 Tabla von Sperling
COSTOS_RANGOS = {
    "Sedimentador 1°":      {"const": (15, 30), "om": (0.8, 1.5)},
    "Laguna anaerobia":      {"const": (12, 30), "om": (0.8, 1.5)},
    "Reactor UASB":          {"const": (20, 35), "om": (1.5, 3.0)},
    "Laguna facultativa":    {"const": (15, 30), "om": (1.0, 2.0)},
    "Lagunas maduración":    {"const": (8, 15),  "om": (0.5, 1.0)},
}

# Tasas de acumulación de lodos (m³/hab·año) — Clase 09 / UASB
TASAS_LODO = {
    "Sedimentador 1°":     0.06,
    "Laguna anaerobia":    0.05,
    "Reactor UASB":        0.03,
    "Laguna facultativa":  0.04,
    "Laguna maduración":   0.015,
}


# ============================================================
# 3. FUNCIONES DE INGENIERÍA
# ============================================================
def proyeccion_geometrica(p0, r, n):
    return p0 * (1 + r / 100) ** n


def calcular_caudales(pob, dot, fr, k1, k2):
    Qmd = pob * dot * fr / 1000
    QMD = k1 * Qmd
    QMH = k2 * QMD
    return Qmd, QMD, QMH


def temperatura_agua(T_aire):
    return 12.7 + 0.54 * T_aire


def carga_volumetrica_anaerobia(T):
    if T < 10:
        return 0.10
    elif T <= 20:
        return 0.02 * T - 0.10
    elif T <= 25:
        return 0.01 * T + 0.10
    else:
        return 0.35


def eficiencia_anaerobia_dbo(T):
    if T <= 25:
        return 2 * T + 20
    return 70.0


def carga_superficial_facultativa(T):
    return 350 * (1.107 - 0.002 * T) ** (T - 25)


def k_corregido(k20, T, theta):
    return k20 * (theta ** (T - 20))


def s_mezcla_completa(So, k, TRH, n=1):
    return So / ((1 + k * TRH / n) ** n)


def s_flujo_piston(So, k, TRH):
    return So * math.exp(-k * TRH)


def s_flujo_disperso(So, k, TRH, d):
    a = math.sqrt(1 + 4 * k * TRH * d)
    num = So * 4 * a * math.exp(1 / (2 * d))
    den = (1 + a) ** 2 * math.exp(a / (2 * d)) - (1 - a) ** 2 * math.exp(-a / (2 * d))
    return num / den


def numero_dispersion(L_W):
    return 1 / L_W


def remover(So, k, TRH, regimen, n=1, L_W=2.0):
    if regimen == "Mezcla completa":
        return s_mezcla_completa(So, k, TRH, n)
    if regimen == "Flujo pistón":
        return s_flujo_piston(So, k, TRH * n)
    d = numero_dispersion(L_W)
    s = So
    for _ in range(int(n)):
        s = s_flujo_disperso(s, k, TRH, d)
    return s


def dimensiones_rectangulares(area, L_W):
    if area <= 0 or L_W <= 0:
        return 0.0, 0.0
    W = math.sqrt(area / L_W)
    L = L_W * W
    return L, W


def diametro_circular(area):
    if area <= 0:
        return 0.0
    return math.sqrt(4 * area / math.pi)


def trh_uasb_por_temperatura(T):
    if T < 16:
        return 14.0
    elif T < 20:
        return 12.0
    elif T <= 26:
        return 8.0
    else:
        return 6.0


def K_metano(T, P=1.0):
    R = 0.08206
    K = 64
    return P * K / (R * (273 + T))


# ----- LODOS -----
def calcular_lodos(unidad_nombre, pob, A_unidad, H_util,
                    densidad=1100, pct_sst=0.05, carga_lecho=150):
    """Manejo de lodos: V/año, frecuencia, área lechos."""
    tasa = TASAS_LODO.get(unidad_nombre, 0.05)
    V_anual = pob * tasa
    altura_anual = V_anual / A_unidad if A_unidad > 0 else 0
    altura_max = H_util / 3
    anos_evac = (altura_max / altura_anual) if altura_anual > 0 else float("inf")
    masa_sst = V_anual * densidad * pct_sst
    A_lecho = masa_sst / carga_lecho if carga_lecho > 0 else 0
    return {
        "tasa": tasa,
        "V_anual": V_anual,
        "altura_anual_cm": altura_anual * 100,
        "anos_evacuacion": anos_evac,
        "masa_sst_anual": masa_sst,
        "A_lecho": A_lecho,
    }


# ----- COSTOS -----
def estimar_costos(opcion, pob, anios=25):
    """Estima costos según rangos US$/hab (Clase 09)."""
    if opcion not in COSTOS_RANGOS:
        return None
    r = COSTOS_RANGOS[opcion]
    c_med = (r["const"][0] + r["const"][1]) / 2
    om_med = (r["om"][0] + r["om"][1]) / 2
    return {
        "construccion_min": r["const"][0] * pob,
        "construccion_max": r["const"][1] * pob,
        "construccion_med": c_med * pob,
        "om_anual_min": r["om"][0] * pob,
        "om_anual_max": r["om"][1] * pob,
        "om_anual_med": om_med * pob,
        "total_25_med": c_med * pob + om_med * pob * anios,
    }


# ----- ENERGÍA DEL BIOGÁS -----
def energia_biogas(Q_CH4_d, eficiencia=0.35, PCI=9.97):
    """kWh/d a partir de m³ CH4/d. PCI ~9,97 kWh/m³ CH4."""
    energia_termica_d = Q_CH4_d * PCI
    energia_electrica_d = energia_termica_d * eficiencia
    return {
        "PCI": PCI,
        "energia_termica_d": energia_termica_d,
        "energia_electrica_d": energia_electrica_d,
        "MWh_anio": energia_electrica_d * 365 / 1000,
        "potencia_kw": energia_electrica_d / 24,
    }


# ----- CUMPLIMIENTO RES 0631 -----
def verificar_res0631(DBO, DQO, SST, Coli):
    items = [
        ("DBO₅ ≤ 90 mg/L (Res 0631)", DBO, RES_0631["DBO5_max"], "mg/L"),
        ("DQO ≤ 180 mg/L (Res 0631)", DQO, RES_0631["DQO_max"], "mg/L"),
        ("SST ≤ 90 mg/L (Res 0631)", SST, RES_0631["SST_max"], "mg/L"),
        ("Coliformes ≤ 1.000 (riego restringido, Dec. 1076)",
            Coli, REUSO_LIMITES["Riego restringido"], "NMP/100mL"),
        ("Coliformes ≤ 100 (riego sin restricción, Dec. 1076)",
            Coli, REUSO_LIMITES["Riego no restringido"], "NMP/100mL"),
    ]
    rows = []
    for nombre, valor, limite, unidad in items:
        cumple = valor <= limite
        rows.append({
            "Parámetro": nombre,
            "Valor": valor,
            "Límite": limite,
            "Unidad": unidad,
            "Cumple": "✓" if cumple else "✗",
        })
    return rows


# ============================================================
# 4. FUNCIONES DE DISEÑO
# ============================================================
def disenar_sedimentador_primario(Qmd, QMH, DBO, DQO, SST, Coli,
                                    TDS_med=40, TDS_pico=100,
                                    H=3.0, LW=2.0, n_paralelo=1,
                                    ef_dbo=35, ef_dqo=35, ef_sst=60, ef_coli=30):
    A_med = Qmd / TDS_med
    A_pico = QMH / TDS_pico
    A = max(A_med, A_pico)
    A_unit = A / n_paralelo
    V = A * H
    V_unit = V / n_paralelo
    TRH_h = V * 24 / Qmd
    L_unit, W_unit = dimensiones_rectangulares(A_unit, LW)
    return {
        "tipo": "Sedimentador 1°",
        "A": A, "A_unit": A_unit, "A_med": A_med, "A_pico": A_pico,
        "V": V, "V_unit": V_unit, "H": H, "L": L_unit, "W": W_unit,
        "n_paralelo": n_paralelo,
        "TRH_h": TRH_h, "TDS_med": TDS_med, "TDS_pico": TDS_pico,
        "DBO_out": DBO * (1 - ef_dbo / 100),
        "DQO_out": DQO * (1 - ef_dqo / 100),
        "SST_out": SST * (1 - ef_sst / 100),
        "Coli_out": Coli * (1 - ef_coli / 100),
        "ef_dbo": ef_dbo, "ef_dqo": ef_dqo,
        "ef_sst": ef_sst, "ef_coli": ef_coli,
    }


def disenar_laguna_anaerobia(Q, DBO, DQO, SST, Coli, T,
                                Lv=None, H=4.0, LW=2.0, n_paralelo=1,
                                talud=2.0,
                                ef_dqo=65, ef_sst=70, ef_coli=90):
    if Lv is None:
        Lv = carga_volumetrica_anaerobia(T)
    L_org = Q * DBO / 1000
    V = L_org / Lv
    V_unit = V / n_paralelo
    A = V / H
    A_unit = A / n_paralelo
    L_d, W = dimensiones_rectangulares(A_unit, LW)
    TRH_d = V / Q
    ef_dbo = eficiencia_anaerobia_dbo(T)
    return {
        "tipo": "Laguna anaerobia",
        "Lv": Lv, "L_org": L_org,
        "A": A, "A_unit": A_unit, "V": V, "V_unit": V_unit,
        "H": H, "L": L_d, "W": W, "talud": talud, "n_paralelo": n_paralelo,
        "TRH_d": TRH_d, "TRH_h": TRH_d * 24,
        "DBO_out": DBO * (1 - ef_dbo / 100),
        "DQO_out": DQO * (1 - ef_dqo / 100),
        "SST_out": SST * (1 - ef_sst / 100),
        "Coli_out": Coli * (1 - ef_coli / 100),
        "ef_dbo": ef_dbo, "ef_dqo": ef_dqo,
        "ef_sst": ef_sst, "ef_coli": ef_coli,
    }


def disenar_reactor_uasb(Q, QMH, DBO, DQO, SST, Coli, T,
                           TRH_h=None, H=5.0, LW=1.0,
                           forma="Circular", n_paralelo=1,
                           Y_acid=0.15, Y_metano=0.03,
                           ef_dbo=70, ef_dqo=70, ef_sst=70, ef_coli=90):
    if TRH_h is None:
        TRH_h = trh_uasb_por_temperatura(T)
    TRH_d = TRH_h / 24
    V = Q * TRH_d
    V_unit = V / n_paralelo
    A = V / H
    A_unit = A / n_paralelo

    L_d, W = dimensiones_rectangulares(A_unit, LW)
    D = diametro_circular(A_unit)

    Vs_md = H / TRH_d
    Vs_mh = Vs_md / 24
    # Verificación con QMH
    Vs_mh_QMH = (QMH * H / V) / 24    # m/h con caudal pico
    Lh = 1 / TRH_d
    Lo = (DQO / 1000) / TRH_d

    DBO_out = DBO * (1 - ef_dbo / 100)
    DQO_out = DQO * (1 - ef_dqo / 100)
    SST_out = SST * (1 - ef_sst / 100)
    Coli_out = Coli * (1 - ef_coli / 100)

    DQO_rem = (DQO - DQO_out) * Q / 1000
    DQO_disp = DQO_rem * (1 - Y_acid)
    DQO_metano = DQO_disp * (1 - Y_metano)
    K_T = K_metano(T)
    Q_CH4 = DQO_metano / K_T if K_T > 0 else 0.0

    return {
        "tipo": "Reactor UASB",
        "TRH_h": TRH_h, "TRH_d": TRH_d,
        "V": V, "V_unit": V_unit, "A": A, "A_unit": A_unit,
        "H": H, "L": L_d, "W": W, "D": D, "forma": forma,
        "n_paralelo": n_paralelo,
        "Vs_mh": Vs_mh, "Vs_mh_QMH": Vs_mh_QMH,
        "Lh": Lh, "Lo": Lo,
        "DBO_out": DBO_out, "DQO_out": DQO_out,
        "SST_out": SST_out, "Coli_out": Coli_out,
        "ef_dbo": ef_dbo, "ef_dqo": ef_dqo,
        "ef_sst": ef_sst, "ef_coli": ef_coli,
        "DQO_rem": DQO_rem, "DQO_metano": DQO_metano,
        "Q_CH4": Q_CH4, "K_T": K_T,
        "Y_acid": Y_acid, "Y_metano": Y_metano,
    }


def recomendar_primario(area_disp, sed, ana, uasb):
    viable_sed = sed["A"] <= area_disp
    viable_ana = ana["A"] <= area_disp
    viable_uasb = uasb["A"] <= area_disp
    viables = []
    if viable_sed: viables.append("Sedimentador 1°")
    if viable_ana: viables.append("Laguna anaerobia")
    if viable_uasb: viables.append("Reactor UASB")

    if not viables:
        return ("Ninguna opción cabe", "danger",
                "El área disponible es insuficiente. Considera aumentar el área "
                "o usar reactores UASB modulares.", viables)

    if viable_uasb:
        rec = "Reactor UASB"
        just = ("UASB es la opción más recomendable: alta eficiencia (65-75%), "
                "huella compacta, recuperación de biogás. Considera arranque ~6 meses.")
    elif viable_ana:
        rec = "Laguna anaerobia"
        just = ("Laguna anaerobia viable, eficiencia DBO según T. Robusta, "
                "pero requiere más área y puede generar olores.")
    else:
        rec = "Sedimentador 1°"
        just = ("Solo el sedimentador cabe. Eficiencia DBO baja (30-40%); "
                "requiere secundario robusto.")
    return rec, "success", just, viables


# ============================================================
# 5. PLANOS ESQUEMÁTICOS (matplotlib)
# ============================================================
COLOR_AGUA = "#cfe8ff"
COLOR_LODO = "#8b6f47"
COLOR_CONCRETO = "#bbbbbb"
COLOR_BIOGAS = "#fff2b3"
COLOR_BORDE = "#222222"


def _setup_axis(ax, title="", xlim=None, ylim=None):
    ax.set_aspect("equal")
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=10, weight="bold")
    if xlim: ax.set_xlim(xlim)
    if ylim: ax.set_ylim(ylim)


def dibujar_sedimentador(L, W, H, n=1, fig_size=(11, 4.5)):
    """Plano + corte de sedimentador primario rectangular."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=fig_size)

    # ---- PLANTA ----
    ax1.add_patch(Rectangle((0, 0), L, W, facecolor=COLOR_AGUA,
                              edgecolor=COLOR_BORDE, linewidth=1.5))
    # Inlet
    ax1.add_patch(Rectangle((-L*0.04, W*0.4), L*0.04, W*0.2,
                              facecolor=COLOR_CONCRETO, edgecolor="black"))
    ax1.annotate("Q", xy=(0, W/2), xytext=(-L*0.13, W/2),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=11, color="blue", va="center", ha="right")
    # Bafle
    ax1.plot([L*0.85]*2, [0, W], "k-", lw=1.2)
    ax1.text(L*0.85, W*1.04, "Bafle", ha="center", fontsize=8)
    # Vertedero
    ax1.plot([L]*2, [0, W], "k-", lw=2.5)
    ax1.text(L, W*1.04, "Vertedero", ha="center", fontsize=8)
    ax1.annotate("Q", xy=(L*1.13, W/2), xytext=(L, W/2),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=11, color="blue", va="center")
    # Tolva
    ax1.add_patch(Rectangle((L*0.05, W*0.05), L*0.78, W*0.1,
                              facecolor=COLOR_LODO, alpha=0.4))
    ax1.text(L/2, W*0.1, "Tolva de lodos", ha="center", va="center", fontsize=7)
    # Dim
    ax1.annotate("", xy=(L, -W*0.18), xytext=(0, -W*0.18),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax1.text(L/2, -W*0.27, f"L = {L:.1f} m", ha="center", fontsize=9, weight="bold")
    ax1.annotate("", xy=(-L*0.08, W), xytext=(-L*0.08, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax1.text(-L*0.16, W/2, f"W = {W:.1f} m", va="center", rotation=90, fontsize=9, weight="bold")
    title_pl = f"PLANTA — {n} módulo(s) en paralelo" if n > 1 else "PLANTA"
    _setup_axis(ax1, title_pl, xlim=(-L*0.27, L*1.27), ylim=(-W*0.45, W*1.2))

    # ---- CORTE ----
    pts_tank = [[0, 0], [0, H], [L, H], [L, H*0.05], [L*0.5, 0]]
    ax2.add_patch(Polygon(pts_tank, facecolor=COLOR_AGUA,
                            edgecolor=COLOR_BORDE, linewidth=1.5))
    # Nivel agua
    ax2.plot([0, L], [H*0.92]*2, "b--", lw=1, alpha=0.6)
    # Lodos
    pts_lodo = [[L*0.05, 0], [L*0.55, 0], [L*0.95, H*0.13], [L*0.05, H*0.13]]
    ax2.add_patch(Polygon(pts_lodo, facecolor=COLOR_LODO, alpha=0.7))
    ax2.text(L*0.4, H*0.07, "Lodos", ha="center", va="center", fontsize=8, color="white")
    # Bafle
    ax2.plot([L*0.85]*2, [H*0.4, H*0.95], "k-", lw=2)
    # Q in
    ax2.annotate("Q", xy=(0, H*0.6), xytext=(-L*0.1, H*0.6),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=11, color="blue", va="center")
    # Q out
    ax2.annotate("Q", xy=(L*1.1, H*0.85), xytext=(L, H*0.85),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=11, color="blue", va="center")
    # Lodos out
    ax2.annotate("Lodos", xy=(L*0.5, -H*0.18), xytext=(L*0.5, 0),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="brown"),
                  fontsize=8, color="brown", ha="center")
    # Dim
    ax2.annotate("", xy=(-L*0.05, H), xytext=(-L*0.05, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax2.text(-L*0.13, H/2, f"H = {H:.1f} m", va="center", rotation=90, fontsize=9, weight="bold")
    ax2.annotate("", xy=(L, -H*0.4), xytext=(0, -H*0.4),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax2.text(L/2, -H*0.5, f"L = {L:.1f} m", ha="center", fontsize=9, weight="bold")
    ax2.text(L*0.97, H*0.97, "▽", fontsize=10, color="blue")
    _setup_axis(ax2, "CORTE LONGITUDINAL",
                  xlim=(-L*0.22, L*1.22), ylim=(-H*0.7, H*1.2))

    plt.tight_layout()
    return fig


def dibujar_laguna(L, W, H, talud=2, tipo="Anaerobia",
                     color_agua=COLOR_AGUA, fig_size=(11, 4.5)):
    """Plano + corte de laguna trapezoidal."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=fig_size)

    # Bottom dimensions (account for talud reducing area)
    L_bot = max(L - 2 * H * talud, L * 0.1)
    W_bot = max(W - 2 * H * talud, W * 0.1)

    # ---- PLANTA ----
    ax1.add_patch(Rectangle((0, 0), L, W, facecolor=color_agua,
                              edgecolor=COLOR_BORDE, linewidth=1.5))
    # Bottom outline
    ox, oy = (L - L_bot)/2, (W - W_bot)/2
    ax1.add_patch(Rectangle((ox, oy), L_bot, W_bot,
                              fill=False, edgecolor="gray",
                              linewidth=1, linestyle="--"))
    ax1.text(L/2, W*0.5, "Espejo de agua", ha="center", va="center",
              fontsize=8, color="#555", style="italic")
    # Inlet
    ax1.plot(0, W*0.5, "o", color="blue", markersize=10)
    ax1.annotate("Q", xy=(0, W*0.5), xytext=(-L*0.1, W*0.5),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=11, color="blue", va="center")
    # Outlet
    ax1.plot(L, W*0.5, "s", color="red", markersize=8)
    ax1.annotate("Q", xy=(L*1.1, W*0.5), xytext=(L, W*0.5),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=11, color="blue", va="center")
    # Dim
    ax1.annotate("", xy=(L, -W*0.18), xytext=(0, -W*0.18),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax1.text(L/2, -W*0.27, f"L (corona) = {L:.1f} m",
              ha="center", fontsize=9, weight="bold")
    ax1.annotate("", xy=(-L*0.08, W), xytext=(-L*0.08, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax1.text(-L*0.16, W/2, f"W = {W:.1f} m",
              va="center", rotation=90, fontsize=9, weight="bold")
    ax1.text(L/2, oy*0.4 if oy > 0 else -W*0.05,
              f"L_fondo = {L_bot:.1f} m",
              ha="center", fontsize=8, color="gray", style="italic")
    _setup_axis(ax1, "PLANTA", xlim=(-L*0.25, L*1.25), ylim=(-W*0.45, W*1.2))

    # ---- CORTE TRANSVERSAL ----
    pts = [[0, 0], [(W - W_bot)/2, H], [W - (W - W_bot)/2, H], [W, 0]]
    ax2.add_patch(Polygon(pts, facecolor=color_agua,
                            edgecolor=COLOR_BORDE, linewidth=1.5))
    # Nivel
    bx_l = (W - W_bot)/2 * 1.05
    bx_r = W - (W - W_bot)/2 * 1.05
    ax2.plot([bx_l, bx_r], [H*0.93]*2, "b--", lw=1, alpha=0.6)
    ax2.text(W*0.05, H*1.0, "↓ Borde libre", fontsize=8, color="gray")
    # Lodos
    sludge_h = H * 0.12
    inset = sludge_h * talud
    pts_lodo = [
        [(W - W_bot)/2, 0],
        [W - (W - W_bot)/2, 0],
        [W - (W - W_bot)/2 - inset, sludge_h],
        [(W - W_bot)/2 + inset, sludge_h],
    ]
    ax2.add_patch(Polygon(pts_lodo, facecolor=COLOR_LODO, alpha=0.7))
    ax2.text(W/2, sludge_h*0.4, "Lodos digeridos",
              ha="center", va="center", fontsize=8, color="white")
    # Talud
    ax2.text((W - W_bot)/4, H*0.5, f"1\n―\n{talud:.0f}",
              ha="center", va="center", fontsize=8,
              bbox=dict(boxstyle="round", facecolor="white", alpha=0.7))
    # Suelo
    ax2.plot([-W*0.05, W*1.05], [0]*2, "k-", lw=0.8)
    # Dim
    ax2.annotate("", xy=(-W*0.08, H), xytext=(-W*0.08, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax2.text(-W*0.15, H/2, f"H = {H:.1f} m",
              va="center", rotation=90, fontsize=9, weight="bold")
    ax2.annotate("", xy=(W, -H*0.3), xytext=(0, -H*0.3),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax2.text(W/2, -H*0.42, f"W (corona) = {W:.1f} m",
              ha="center", fontsize=9, weight="bold")
    # Q in/out
    ax2.annotate("Q", xy=(0, H*0.7), xytext=(-W*0.1, H*0.7),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center")
    ax2.annotate("Q", xy=(W*1.1, H*0.85), xytext=(W, H*0.85),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center")
    _setup_axis(ax2, f"CORTE TRANSVERSAL — talud {talud:.0f}:1",
                  xlim=(-W*0.22, W*1.22), ylim=(-H*0.55, H*1.2))

    plt.suptitle(f"Laguna {tipo}", fontsize=12, weight="bold", y=1.02)
    plt.tight_layout()
    return fig


def dibujar_uasb(forma, L, W, H, D, fig_size=(7, 6)):
    """Corte vertical de reactor UASB (circular o rectangular)."""
    fig, ax = plt.subplots(figsize=fig_size)

    if forma == "Circular":
        ancho_total = D
        labels = (f"D = {D:.1f} m", f"H = {H:.1f} m")
        title = f"REACTOR UASB CIRCULAR — D = {D:.1f} m, H = {H:.1f} m"
    else:
        ancho_total = W
        labels = (f"W = {W:.1f} m", f"H = {H:.1f} m")
        title = f"REACTOR UASB RECTANGULAR — W = {W:.1f} m, H = {H:.1f} m"

    a = ancho_total
    # Cuerpo principal
    ax.add_patch(Rectangle((0, 0), a, H,
                            facecolor=COLOR_AGUA,
                            edgecolor=COLOR_BORDE, linewidth=1.5))
    # Lecho de lodo (denso, fondo)
    ax.add_patch(Rectangle((0, 0), a, H*0.20,
                            facecolor="#5d3a1f", alpha=0.85))
    ax.text(a/2, H*0.10, "Lecho de lodo (40-100 g/L)",
              ha="center", va="center", fontsize=9, color="white", weight="bold")
    # Manto de lodo
    ax.add_patch(Rectangle((0, H*0.20), a, H*0.30,
                            facecolor=COLOR_LODO, alpha=0.55))
    ax.text(a/2, H*0.35, "Manto de lodo (10-30 g/L)",
              ha="center", va="center", fontsize=9, color="black")
    # Separador GLS (trapezoide)
    pts_sep = [[a*0.05, H*0.65], [a*0.30, H*0.85],
                [a*0.70, H*0.85], [a*0.95, H*0.65]]
    ax.add_patch(Polygon(pts_sep, fill=True, facecolor=COLOR_CONCRETO,
                          edgecolor="black", linewidth=1.2))
    ax.text(a/2, H*0.78, "Separador G-L-S",
              ha="center", va="center", fontsize=9, weight="bold")
    # Cámara de biogás
    ax.add_patch(Rectangle((a*0.30, H*0.85), a*0.40, H*0.13,
                            facecolor=COLOR_BIOGAS, edgecolor="black", linewidth=1))
    ax.text(a/2, H*0.92, "Biogás", ha="center", va="center", fontsize=9, weight="bold")
    # Salida biogás
    ax.annotate("CH₄ + CO₂", xy=(a*0.5, H*1.15), xytext=(a*0.5, H*0.98),
                  arrowprops=dict(arrowstyle="->", lw=2, color="orange"),
                  fontsize=9, color="orange", ha="center", va="bottom", weight="bold")
    # Distribuidor influente (fondo)
    n_orif = 5
    for i in range(1, n_orif):
        x = a * i / n_orif
        ax.annotate("", xy=(x, H*0.05), xytext=(x, -H*0.1),
                      arrowprops=dict(arrowstyle="->", color="blue", lw=1))
    ax.text(a/2, -H*0.15, "Influente (distribuidor)",
              ha="center", fontsize=9, color="blue")
    # Salida efluente
    ax.annotate("Efluente", xy=(a*1.18, H*0.78),
                  xytext=(a, H*0.78),
                  arrowprops=dict(arrowstyle="->", lw=2, color="green"),
                  fontsize=9, color="green", va="center", weight="bold")
    # Dim
    ax.annotate("", xy=(-a*0.06, H), xytext=(-a*0.06, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax.text(-a*0.13, H/2, labels[1],
              va="center", rotation=90, fontsize=10, weight="bold")
    ax.annotate("", xy=(a, -H*0.25), xytext=(0, -H*0.25),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax.text(a/2, -H*0.34, labels[0], ha="center", fontsize=10, weight="bold")
    _setup_axis(ax, title,
                  xlim=(-a*0.25, a*1.35), ylim=(-H*0.45, H*1.3))
    plt.tight_layout()
    return fig


def dibujar_lagunas_serie(L, W, H, n, fig_size=(13, 4)):
    """Vista en planta de n lagunas de maduración en serie."""
    fig, ax = plt.subplots(figsize=fig_size)
    sep = L * 0.05
    for i in range(int(n)):
        x = i * (L + sep)
        ax.add_patch(Rectangle((x, 0), L, W, facecolor=COLOR_AGUA,
                                  edgecolor=COLOR_BORDE, linewidth=1.5))
        ax.text(x + L/2, W/2, f"Mad. {i+1}",
                  ha="center", va="center", fontsize=10, weight="bold")
        if i < n - 1:
            ax.annotate("", xy=(x + L + sep, W/2), xytext=(x + L, W/2),
                          arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"))
        else:
            ax.annotate("Q", xy=(x + L + sep, W/2), xytext=(x + L, W/2),
                          arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                          fontsize=11, color="blue", va="center")
    ax.annotate("Q", xy=(0, W/2), xytext=(-sep, W/2),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=11, color="blue", va="center", ha="right")
    # Dim
    ax.annotate("", xy=(L, -W*0.15), xytext=(0, -W*0.15),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax.text(L/2, -W*0.25, f"L = {L:.1f} m", ha="center", fontsize=9)
    ax.annotate("", xy=(-sep, W), xytext=(-sep, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax.text(-sep*1.5, W/2, f"W = {W:.1f} m",
              va="center", rotation=90, fontsize=9)
    total_L = n * L + (n - 1) * sep
    _setup_axis(ax,
                  f"LAGUNAS DE MADURACIÓN — {int(n)} en serie · H = {H:.1f} m",
                  xlim=(-sep*2, total_L + sep*2),
                  ylim=(-W*0.4, W*1.2))
    plt.tight_layout()
    return fig


def dibujar_tren(unidades, fig_size=(13, 4)):
    """Diagrama esquemático del tren de tratamiento."""
    fig, ax = plt.subplots(figsize=fig_size)
    n = len(unidades)
    box_w, box_h = 1.0, 0.7
    sep = 0.6
    y = 0
    colores = ["#a8d8ea", "#aa96da", "#fcbad3", "#ffffd2", "#a0e8af"]
    for i, u in enumerate(unidades):
        x = i * (box_w + sep)
        c = colores[i % len(colores)]
        box = FancyBboxPatch((x, y), box_w, box_h,
                              boxstyle="round,pad=0.02",
                              facecolor=c, edgecolor="black", linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + box_w/2, y + box_h*0.7, u["nombre"],
                  ha="center", fontsize=10, weight="bold")
        ax.text(x + box_w/2, y + box_h*0.45,
                  f"V = {u['V']:,.0f} m³",
                  ha="center", fontsize=8)
        ax.text(x + box_w/2, y + box_h*0.25,
                  f"TRH = {u['TRH']}",
                  ha="center", fontsize=8)
        # Salidas
        ax.text(x + box_w/2, y - 0.15,
                  f"DBO = {u['DBO']:.0f} mg/L",
                  ha="center", fontsize=8, color="darkred")
        ax.text(x + box_w/2, y - 0.30,
                  f"Coli = {u['Coli']:.1e}",
                  ha="center", fontsize=7, color="darkblue")
        # Flecha al siguiente
        if i < n - 1:
            ax.annotate("", xy=(x + box_w + sep, y + box_h/2),
                          xytext=(x + box_w, y + box_h/2),
                          arrowprops=dict(arrowstyle="->", lw=2, color="blue"))
    # Entrada inicial
    ax.annotate("Afluente",
                  xy=(0, y + box_h/2), xytext=(-sep*0.6, y + box_h/2),
                  arrowprops=dict(arrowstyle="->", lw=2, color="blue"),
                  fontsize=9, color="blue", va="center", ha="right")
    # Salida final
    ax.annotate("Efluente",
                  xy=((n-1)*(box_w+sep) + box_w + sep*0.6, y + box_h/2),
                  xytext=((n-1)*(box_w+sep) + box_w, y + box_h/2),
                  arrowprops=dict(arrowstyle="->", lw=2, color="green"),
                  fontsize=9, color="green", va="center")

    _setup_axis(ax, "TREN DE TRATAMIENTO",
                  xlim=(-sep*1.5, n*(box_w + sep) + sep),
                  ylim=(-0.6, box_h + 0.3))
    plt.tight_layout()
    return fig


# ============================================================
# 6. EXPORTACIÓN A WORD
# ============================================================
def fig_to_buf(fig, dpi=120):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    return buf


def generar_word_memoria(contexto):
    """Genera memoria de cálculo en Word con esquemas embebidos."""
    if not DOCX_OK:
        return None

    doc = Document()
    # Encabezado
    doc.add_heading("Memoria de Cálculo - PTAR Barbosa", 0)
    p = doc.add_paragraph()
    p.add_run("Materia: ").bold = True
    p.add_run("Sistemas de Tratamiento de Aguas Residuales · UdeA\n")
    p.add_run("Caso: ").bold = True
    p.add_run(f"{contexto['municipio']}\n")
    p.add_run("Tren seleccionado: ").bold = True
    p.add_run(f"{contexto['primario']} → Facultativa → Maduración")

    # 1. Datos del proyecto
    doc.add_heading("1. Datos del proyecto", 1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    h = t.rows[0].cells
    h[0].text = "Parámetro"
    h[1].text = "Valor"
    for k, v in contexto["datos"].items():
        row = t.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    # 2. Tren de tratamiento (figura)
    doc.add_heading("2. Tren de tratamiento", 1)
    if contexto.get("fig_tren") is not None:
        doc.add_picture(contexto["fig_tren"], width=Inches(6.5))

    # 3. Memorias por unidad
    doc.add_heading("3. Memorias de cálculo por unidad", 1)
    for nombre, mem in contexto["memorias"].items():
        doc.add_heading(nombre, 2)
        if "fig" in mem and mem["fig"] is not None:
            doc.add_picture(mem["fig"], width=Inches(6.0))
        for linea in mem["texto"].split("\n"):
            if linea.strip():
                doc.add_paragraph(linea)

    # 4. Resumen
    doc.add_heading("4. Resumen y eficiencias", 1)
    if "resumen_df" in contexto:
        df = contexto["resumen_df"]
        t2 = doc.add_table(rows=1, cols=len(df.columns))
        t2.style = "Light Grid Accent 1"
        for i, c in enumerate(df.columns):
            t2.rows[0].cells[i].text = str(c)
        for _, row in df.iterrows():
            cells = t2.add_row().cells
            for i, c in enumerate(df.columns):
                cells[i].text = str(row[c])

    # 5. Cumplimiento normativo
    doc.add_heading("5. Verificación Resolución 0631 de 2015", 1)
    if "cumplimiento" in contexto:
        df = pd.DataFrame(contexto["cumplimiento"])
        t3 = doc.add_table(rows=1, cols=len(df.columns))
        t3.style = "Light Grid Accent 1"
        for i, c in enumerate(df.columns):
            t3.rows[0].cells[i].text = str(c)
        for _, row in df.iterrows():
            cells = t3.add_row().cells
            for i, c in enumerate(df.columns):
                cells[i].text = str(row[c])

    # 6. Footer
    doc.add_paragraph()
    doc.add_paragraph(
        "Generado automáticamente. Referencias: Mara (1997), "
        "von Sperling & Chernicharo (2005), van Haandel & Lettinga (1994), "
        "RAS 2017, Res 0631/2015."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 7. SIDEBAR
# ============================================================
st.sidebar.title("⚙️ Parámetros de diseño")
st.sidebar.caption(f"📍 Caso: {DATOS_BARBOSA['municipio']}")

with st.sidebar.expander("📥 Datos del afluente", expanded=False):
    DBO_in = st.number_input("DBO₅ (mg/L)", value=DATOS_BARBOSA["DBO5"],
                              min_value=1.0, step=10.0)
    DQO_in = st.number_input("DQO (mg/L)", value=DATOS_BARBOSA["DQO"],
                              min_value=1.0, step=10.0)
    SST_in = st.number_input("SST (mg/L)", value=DATOS_BARBOSA["SST"],
                              min_value=1.0, step=10.0)
    Coli_in = st.number_input("Coliformes fecales (NMP/100mL)",
                               value=DATOS_BARBOSA["coliformes_fecales"],
                               format="%.2e", min_value=1.0)

with st.sidebar.expander("🌡️ Condiciones ambientales", expanded=True):
    T_aire = st.slider("T aire mes más frío (°C)",
                       min_value=10.0, max_value=30.0,
                       value=DATOS_BARBOSA["T_aire_mes_frio"], step=0.5)
    T_agua_yanez = temperatura_agua(T_aire)
    T_agua = st.slider("T agua del agua residual (°C)",
                        min_value=10.0, max_value=30.0,
                        value=round(T_agua_yanez, 1), step=0.5,
                        help=f"Yáñez: T_agua = {T_agua_yanez:.1f} °C")

with st.sidebar.expander("🌊 Régimen hidráulico", expanded=False):
    regimen = st.selectbox("Modelo hidráulico",
                            ["Mezcla completa", "Flujo pistón", "Flujo disperso"],
                            index=0)

with st.sidebar.expander("🏗️ Tratamiento primario", expanded=True):
    area_disp = st.number_input(
        "📐 Área disponible (m²)",
        value=DATOS_BARBOSA["area_disponible"],
        min_value=0.0, step=500.0,
    )
    tipo_primario = st.radio(
        "Primario activo",
        ["Sedimentador 1°", "Laguna anaerobia", "Reactor UASB"],
        index=2,
    )

with st.sidebar.expander("👥 Demografía", expanded=False):
    pob_actual = st.number_input("Población actual (hab)",
                                   value=DATOS_BARBOSA["poblacion_actual"],
                                   min_value=1, step=500)
    tasa = st.number_input("Tasa crecimiento (%/año)",
                             value=DATOS_BARBOSA["tasa_crecimiento"],
                             min_value=0.0, step=0.1)
    horizonte = st.number_input("Horizonte (años)",
                                  value=DATOS_BARBOSA["horizonte_diseno"],
                                  min_value=1, step=1)
    dotacion = st.number_input("Dotación (L/hab·d)",
                                 value=DATOS_BARBOSA["dotacion_neta"],
                                 min_value=1, step=10)
    factor_r = st.number_input("Factor de retorno",
                                 value=DATOS_BARBOSA["factor_retorno"],
                                 min_value=0.05, max_value=1.0, step=0.05)
    k1_in = st.number_input("k1", value=DATOS_BARBOSA["k1"],
                              min_value=1.0, step=0.05)
    k2_in = st.number_input("k2", value=DATOS_BARBOSA["k2"],
                              min_value=1.0, step=0.05)


# ============================================================
# 8. VALIDACIONES
# ============================================================
errores = []
if DBO_in <= 0 or DQO_in <= 0 or SST_in <= 0:
    errores.append("DBO, DQO y SST deben ser positivas.")
if Coli_in <= 0:
    errores.append("Coliformes deben ser positivos.")
if DBO_in > DQO_in:
    errores.append("DBO no puede ser mayor que DQO.")
if pob_actual <= 0 or dotacion <= 0:
    errores.append("Población y dotación deben ser positivas.")
if not (0 < factor_r <= 1):
    errores.append("Factor de retorno entre 0 y 1.")
if k1_in < 1 or k2_in < 1:
    errores.append("k1 y k2 ≥ 1.")
if T_agua <= 0:
    errores.append("Temperatura del agua > 0.")
if area_disp < 0:
    errores.append("Área disponible ≥ 0.")

if errores:
    st.error("❌ Errores en datos de entrada:")
    for er in errores:
        st.error(er)
    st.stop()


# ============================================================
# 9. CÁLCULOS PRELIMINARES
# ============================================================
pob_diseno = proyeccion_geometrica(pob_actual, tasa, horizonte)
Qmd, QMD, QMH = calcular_caudales(pob_diseno, dotacion, factor_r, k1_in, k2_in)
Q_d = Qmd

# Diseños preliminares (defaults para selección)
sed_pre = disenar_sedimentador_primario(Qmd, QMH, DBO_in, DQO_in, SST_in, Coli_in)
ana_pre = disenar_laguna_anaerobia(Q_d, DBO_in, DQO_in, SST_in, Coli_in, T_agua)
uasb_pre = disenar_reactor_uasb(Q_d, QMH, DBO_in, DQO_in, SST_in, Coli_in, T_agua)


# ============================================================
# 10. ENCABEZADO
# ============================================================
st.title("💧 Diseño de PTAR — Caso Barbosa, Antioquia")
st.markdown(f"### Tren: **{tipo_primario}** → Facultativa → Maduración")

c1, c2, c3, c4, c5 = st.columns(5)
c1.metric("Población", f"{pob_diseno:,.0f} hab")
c2.metric("Qmd", f"{Qmd:,.1f} m³/d")
c3.metric("QMD", f"{QMD:,.1f} m³/d")
c4.metric("QMH", f"{QMH:,.1f} m³/d")
c5.metric("Área disp.", f"{area_disp:,.0f} m²")


# ============================================================
# 11. PESTAÑAS
# ============================================================
tabs = st.tabs([
    "📍 Caudales",
    "📐 Selección + Costos",
    "🏛️ Sedimentador",
    "🌊 Anaerobia",
    "⚗️ UASB",
    "🌱 Facultativa",
    "💧 Maduración",
    "🪣 Lodos",
    "📈 Sensibilidad",
    "📊 Resumen",
])
(tab0, tab_sel, tab_sed, tab_ana, tab_uasb,
 tab_fac, tab_mad, tab_lod, tab_sens, tab_res) = tabs


# -----------------------------------------------------------------
# TAB 0 — Caudales
# -----------------------------------------------------------------
with tab0:
    st.header("Proyección poblacional y caudales")
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Proyección geométrica")
        st.latex(r"P_f = P_0 \cdot (1 + r/100)^n")
        df_pob = pd.DataFrame({
            "Año": [DATOS_BARBOSA['anio_actual'],
                    DATOS_BARBOSA['anio_actual'] + horizonte],
            "Población (hab)": [int(pob_actual), int(round(pob_diseno))],
        })
        st.dataframe(df_pob, hide_index=True, use_container_width=True)
    with col2:
        st.subheader("Caudales (RAS 2017)")
        df_q = pd.DataFrame({
            "Caudal": ["Qmd", "QMD", "QMH"],
            "m³/d": [f"{Qmd:,.1f}", f"{QMD:,.1f}", f"{QMH:,.1f}"],
            "L/s": [f"{Qmd*1000/86400:,.1f}",
                    f"{QMD*1000/86400:,.1f}",
                    f"{QMH*1000/86400:,.1f}"],
        })
        st.dataframe(df_q, hide_index=True, use_container_width=True)

    st.info("ℹ️ Lagunas y UASB se diseñan con Qmd. Sedimentador 1° se "
            "diseña con QMH para verificación (RAS 2017 Art. 189).")


# -----------------------------------------------------------------
# TAB SELECCIÓN + COSTOS
# -----------------------------------------------------------------
with tab_sel:
    st.header("📐 Selección de tratamiento primario")

    rec, status, just, viables = recomendar_primario(area_disp, sed_pre, ana_pre, uasb_pre)

    df_comp = pd.DataFrame({
        "Alternativa": ["Sedimentador 1°", "Laguna anaerobia", "Reactor UASB"],
        "Área req (m²)": [sed_pre["A"], ana_pre["A"], uasb_pre["A"]],
        "Volumen (m³)": [sed_pre["V"], ana_pre["V"], uasb_pre["V"]],
        "TRH": [f"{sed_pre['TRH_h']:.1f} h",
                 f"{ana_pre['TRH_d']:.1f} d",
                 f"{uasb_pre['TRH_h']:.1f} h"],
        "Ef. DBO (%)": [sed_pre["ef_dbo"], ana_pre["ef_dbo"], uasb_pre["ef_dbo"]],
        "Ef. DQO (%)": [sed_pre["ef_dqo"], ana_pre["ef_dqo"], uasb_pre["ef_dqo"]],
        "¿Cabe?": [
            "✅" if sed_pre["A"] <= area_disp else "❌",
            "✅" if ana_pre["A"] <= area_disp else "❌",
            "✅" if uasb_pre["A"] <= area_disp else "❌",
        ],
    })
    st.dataframe(
        df_comp.style.format({
            "Área req (m²)": "{:,.0f}",
            "Volumen (m³)": "{:,.0f}",
        }),
        hide_index=True, use_container_width=True,
    )

    if status == "danger":
        st.error(f"⚠️ {rec}")
    else:
        st.success(f"🎯 **Recomendación:** {rec}")
        st.info(just)

    if rec != tipo_primario and status != "danger":
        st.warning(f"💡 Has seleccionado **{tipo_primario}**, pero la "
                   f"recomendación es **{rec}**.")

    # Comparación visual
    st.subheader("Comparación visual de áreas")
    df_areas = pd.DataFrame({
        "Alternativa": ["Sedimentador", "Anaerobia", "UASB", "Disponible"],
        "Área (m²)": [sed_pre["A"], ana_pre["A"], uasb_pre["A"], area_disp],
    })
    st.bar_chart(df_areas.set_index("Alternativa"))

    # Costos
    st.subheader("💰 Costos comparativos (US$/hab — Clase 09)")
    cs = estimar_costos("Sedimentador 1°", pob_diseno)
    ca = estimar_costos("Laguna anaerobia", pob_diseno)
    cu = estimar_costos("Reactor UASB", pob_diseno)
    cf = estimar_costos("Laguna facultativa", pob_diseno)
    cm = estimar_costos("Lagunas maduración", pob_diseno)

    df_costos = pd.DataFrame({
        "Unidad": ["Sedimentador 1°", "Laguna anaerobia", "Reactor UASB",
                    "Laguna facultativa", "Lagunas maduración"],
        "Constr. mín (US$)": [cs["construccion_min"], ca["construccion_min"],
                                cu["construccion_min"], cf["construccion_min"],
                                cm["construccion_min"]],
        "Constr. máx (US$)": [cs["construccion_max"], ca["construccion_max"],
                                cu["construccion_max"], cf["construccion_max"],
                                cm["construccion_max"]],
        "O&M anual (US$)": [cs["om_anual_med"], ca["om_anual_med"],
                              cu["om_anual_med"], cf["om_anual_med"],
                              cm["om_anual_med"]],
        "Total 25 años (US$)": [cs["total_25_med"], ca["total_25_med"],
                                  cu["total_25_med"], cf["total_25_med"],
                                  cm["total_25_med"]],
    })
    st.dataframe(
        df_costos.style.format({c: "{:,.0f}" for c in df_costos.columns
                                  if c != "Unidad"}),
        hide_index=True, use_container_width=True,
    )

    # Costo del tren completo según selección
    st.subheader("💵 Costo del tren seleccionado")
    primario_costo = {"Sedimentador 1°": cs, "Laguna anaerobia": ca,
                       "Reactor UASB": cu}[tipo_primario]
    total_tren_const = (primario_costo["construccion_med"] +
                         cf["construccion_med"] + cm["construccion_med"])
    total_tren_om_anual = (primario_costo["om_anual_med"] +
                            cf["om_anual_med"] + cm["om_anual_med"])
    total_tren_25 = total_tren_const + total_tren_om_anual * 25

    cc1, cc2, cc3 = st.columns(3)
    cc1.metric("Construcción", f"US$ {total_tren_const:,.0f}")
    cc2.metric("O&M anual", f"US$ {total_tren_om_anual:,.0f}")
    cc3.metric("Total 25 años", f"US$ {total_tren_25:,.0f}")


# -----------------------------------------------------------------
# TAB SEDIMENTADOR
# -----------------------------------------------------------------
with tab_sed:
    st.header("🏛️ Sedimentador Primario")
    st.caption("RAS 2017 Art. 189 / Qasim (1999)")

    col_a, col_b = st.columns([1, 2])
    with col_a:
        TDS_med = st.slider("TDS Qmd (m³/m²·d)", 30, 50, 40, 1, key="tds_med")
        TDS_pico = st.slider("TDS QMH (m³/m²·d)", 70, 130, 100, 5, key="tds_pico")
        H_sed = st.slider("Profundidad H (m)", 2.5, 4.0, 3.0, 0.1, key="h_sed")
        LW_sed = st.slider("Relación L/W", 1.5, 15.0, 2.5, 0.5, key="lw_sed")
        n_sed = st.slider("Unidades en paralelo", 1, 4, 2, 1, key="n_sed",
                          help="RAS recomienda mínimo 2 para Q ≥ 100 L/s")

    sed = disenar_sedimentador_primario(
        Qmd, QMH, DBO_in, DQO_in, SST_in, Coli_in,
        TDS_med=TDS_med, TDS_pico=TDS_pico, H=H_sed, LW=LW_sed,
        n_paralelo=n_sed,
    )

    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("Área total", f"{sed['A']:,.0f} m²")
        m2.metric("Área/módulo", f"{sed['A_unit']:,.0f} m²")
        m3.metric("TRH", f"{sed['TRH_h']:.2f} h")
        m4, m5, m6 = st.columns(3)
        m4.metric("L (módulo)", f"{sed['L']:.1f} m")
        m5.metric("W (módulo)", f"{sed['W']:.1f} m")
        m6.metric("V total", f"{sed['V']:,.0f} m³")

        if 1.5 <= sed["TRH_h"] <= 2.5:
            st.success(f"✓ TRH = {sed['TRH_h']:.2f} h (RAS 1.5-2.5 h)")
        else:
            st.warning(f"⚠ TRH = {sed['TRH_h']:.2f} h fuera de rango")
        if sed["A"] <= area_disp:
            st.success(f"✓ Cabe ({sed['A']/area_disp*100:.1f}%)")
        else:
            st.error(f"✗ NO cabe (necesita {sed['A']:,.0f} m²)")

    # Plano esquemático
    st.subheader("📐 Plano esquemático")
    fig_sed = dibujar_sedimentador(sed["L"], sed["W"], sed["H"], sed["n_paralelo"])
    st.pyplot(fig_sed)
    plt.close(fig_sed)

    df_ef_sed = pd.DataFrame({
        "Parámetro": ["DBO₅", "DQO", "SST", "Coliformes"],
        "Entrada": [f"{DBO_in:.1f} mg/L", f"{DQO_in:.1f} mg/L",
                      f"{SST_in:.1f} mg/L", f"{Coli_in:.2e}"],
        "Eficiencia": [f"{sed['ef_dbo']}%", f"{sed['ef_dqo']}%",
                         f"{sed['ef_sst']}%", f"{sed['ef_coli']}%"],
        "Salida": [f"{sed['DBO_out']:.1f} mg/L",
                     f"{sed['DQO_out']:.1f} mg/L",
                     f"{sed['SST_out']:.1f} mg/L",
                     f"{sed['Coli_out']:.2e}"],
    })
    st.dataframe(df_ef_sed, hide_index=True, use_container_width=True)

    with st.expander("📋 Memoria de cálculo"):
        st.markdown(f"""
- A_med = Qmd/TDS = {Qmd:,.0f}/{TDS_med} = **{sed['A_med']:,.1f} m²**
- A_pico = QMH/TDS_pico = {QMH:,.0f}/{TDS_pico} = **{sed['A_pico']:,.1f} m²**
- A (rige) = **{sed['A']:,.0f} m²** (en {sed['n_paralelo']} módulos)
- V = A·H = **{sed['V']:,.0f} m³**
- TRH = V·24/Q = **{sed['TRH_h']:.2f} h** (rango 1.5-2.5 h)
- Por módulo: L = {sed['L']:.1f} m · W = {sed['W']:.1f} m · A = {sed['A_unit']:,.0f} m²
        """)


# -----------------------------------------------------------------
# TAB ANAEROBIA
# -----------------------------------------------------------------
with tab_ana:
    st.header("🌊 Laguna Anaerobia")
    st.caption("Mara (1997) / Clase 09")

    col_a, col_b = st.columns([1, 2])
    with col_a:
        Lv_calc = carga_volumetrica_anaerobia(T_agua)
        st.metric("Lv (Mara)", f"{Lv_calc:.3f} kg/m³·d")
        Lv_ana = st.number_input("Lv adoptado", value=round(Lv_calc, 3),
                                   min_value=0.05, max_value=0.50, step=0.01,
                                   key="lv_ana")
        H_ana = st.slider("H (m)", 3.0, 5.0, 4.0, 0.1, key="h_ana")
        LW_ana = st.slider("L/W", 1.0, 3.0, 2.0, 0.1, key="lw_ana")
        talud_ana = st.slider("Talud m:1", 1.5, 3.0, 2.0, 0.1, key="t_ana")
        n_ana = st.slider("Lagunas en paralelo", 1, 4, 2, 1, key="n_ana")

    ana = disenar_laguna_anaerobia(
        Q_d, DBO_in, DQO_in, SST_in, Coli_in, T_agua,
        Lv=Lv_ana, H=H_ana, LW=LW_ana, talud=talud_ana, n_paralelo=n_ana,
    )

    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("V total", f"{ana['V']:,.0f} m³")
        m2.metric("Área total", f"{ana['A']:,.0f} m²")
        m3.metric("TRH", f"{ana['TRH_d']:.2f} d")
        m4, m5, m6 = st.columns(3)
        m4.metric("L (módulo)", f"{ana['L']:.1f} m")
        m5.metric("W (módulo)", f"{ana['W']:.1f} m")
        m6.metric("Carga DBO", f"{ana['L_org']:,.0f} kg/d")

        if 3 <= ana["TRH_d"] <= 6:
            st.success(f"✓ TRH = {ana['TRH_d']:.2f} d (3-6 d)")
        else:
            st.warning(f"⚠ TRH = {ana['TRH_d']:.2f} d fuera de rango")
        if ana["A"] <= area_disp:
            st.success(f"✓ Cabe ({ana['A']/area_disp*100:.1f}%)")
        else:
            st.error(f"✗ NO cabe ({ana['A']:,.0f} m² necesarios)")

    st.subheader("📐 Plano esquemático")
    fig_ana = dibujar_laguna(ana["L"], ana["W"], ana["H"],
                                talud=ana["talud"], tipo="Anaerobia",
                                color_agua="#9bb8d6")
    st.pyplot(fig_ana)
    plt.close(fig_ana)

    df_ef_ana = pd.DataFrame({
        "Parámetro": ["DBO₅", "DQO", "SST", "Coliformes"],
        "Entrada": [f"{DBO_in:.1f}", f"{DQO_in:.1f}",
                      f"{SST_in:.1f}", f"{Coli_in:.2e}"],
        "Eficiencia": [f"{ana['ef_dbo']:.0f}%", f"{ana['ef_dqo']}%",
                         f"{ana['ef_sst']}%", f"{ana['ef_coli']}%"],
        "Salida": [f"{ana['DBO_out']:.1f}", f"{ana['DQO_out']:.1f}",
                     f"{ana['SST_out']:.1f}", f"{ana['Coli_out']:.2e}"],
    })
    st.dataframe(df_ef_ana, hide_index=True, use_container_width=True)


# -----------------------------------------------------------------
# TAB UASB
# -----------------------------------------------------------------
with tab_uasb:
    st.header("⚗️ Reactor UASB")
    st.caption("van Haandel & Lettinga (1994) / Clase UASB")

    col_a, col_b = st.columns([1, 2])
    with col_a:
        forma_uasb = st.radio("Forma", ["Circular", "Rectangular"], index=0)
        TRH_calc = trh_uasb_por_temperatura(T_agua)
        st.metric("TRH recomendado", f"{TRH_calc:.1f} h")
        TRH_uasb = st.slider("TRH (h)", 4.0, 14.0, TRH_calc, 0.5, key="trh_uasb")
        H_uasb = st.slider("H (m)", 4.0, 6.0, 5.0, 0.1, key="h_uasb")
        LW_uasb = st.slider("L/W (si rectangular)", 1.0, 3.0, 1.0, 0.1, key="lw_u")
        n_uasb = st.slider("Reactores en paralelo", 1, 6, 2, 1, key="n_uasb")
        Y_acid = st.number_input("Y acidogénica", value=0.15, step=0.01, key="ya")
        Y_metano = st.number_input("Y metanogénica", value=0.03, step=0.01, key="ym")

    uasb = disenar_reactor_uasb(
        Q_d, QMH, DBO_in, DQO_in, SST_in, Coli_in, T_agua,
        TRH_h=TRH_uasb, H=H_uasb, LW=LW_uasb,
        forma=forma_uasb, n_paralelo=n_uasb,
        Y_acid=Y_acid, Y_metano=Y_metano,
    )

    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("V total", f"{uasb['V']:,.0f} m³")
        m2.metric("V/módulo", f"{uasb['V_unit']:,.0f} m³")
        m3.metric("TRH", f"{uasb['TRH_h']:.1f} h")
        m4, m5, m6 = st.columns(3)
        m4.metric("Vs (Qmd)", f"{uasb['Vs_mh']:.2f} m/h")
        m5.metric("Vs (QMH)", f"{uasb['Vs_mh_QMH']:.2f} m/h")
        m6.metric("Lo", f"{uasb['Lo']:.2f} kg/m³·d")

        # Verificaciones (Qmd y QMH)
        if uasb["TRH_h"] >= 4:
            st.success(f"✓ TRH = {uasb['TRH_h']:.1f} h ≥ 4 h")
        else:
            st.error(f"✗ TRH = {uasb['TRH_h']:.1f} h < 4 h")

        if 0.5 <= uasb["Vs_mh"] <= 0.7:
            st.success(f"✓ Vs(Qmd) = {uasb['Vs_mh']:.2f} m/h (0.5-0.7)")
        else:
            st.warning(f"⚠ Vs(Qmd) fuera de rango (0.5-0.7)")

        if uasb["Vs_mh_QMH"] <= 1.1:
            st.success(f"✓ Vs(QMH) = {uasb['Vs_mh_QMH']:.2f} m/h ≤ 1.1 (sin lavado)")
        else:
            st.error(f"✗ Vs(QMH) = {uasb['Vs_mh_QMH']:.2f} m/h > 1.1 m/h: riesgo de arrastre")

        if 2.5 <= uasb["Lo"] <= 3.5:
            st.success(f"✓ Lo en rango ARD (2.5-3.5)")
        else:
            st.info(f"ℹ Lo = {uasb['Lo']:.2f} fuera del rango ARD típico")

        if uasb["A"] <= area_disp:
            st.success(f"✓ Cabe ({uasb['A']/area_disp*100:.1f}%)")
        else:
            st.error(f"✗ NO cabe")

    # Plano esquemático
    st.subheader("📐 Plano esquemático (corte vertical)")
    fig_uasb = dibujar_uasb(forma_uasb, uasb["L"], uasb["W"],
                                uasb["H"], uasb["D"])
    st.pyplot(fig_uasb)
    plt.close(fig_uasb)

    if forma_uasb == "Circular":
        st.caption(f"Cada uno de los {n_uasb} reactores: D = {uasb['D']:.1f} m, "
                   f"H = {uasb['H']:.1f} m")
    else:
        st.caption(f"Cada uno de los {n_uasb} reactores: L = {uasb['L']:.1f} m, "
                   f"W = {uasb['W']:.1f} m, H = {uasb['H']:.1f} m")

    # Eficiencias
    df_ef_uasb = pd.DataFrame({
        "Parámetro": ["DBO₅", "DQO", "SST", "Coliformes"],
        "Entrada": [f"{DBO_in:.1f}", f"{DQO_in:.1f}",
                      f"{SST_in:.1f}", f"{Coli_in:.2e}"],
        "Eficiencia": [f"{uasb['ef_dbo']}%", f"{uasb['ef_dqo']}%",
                         f"{uasb['ef_sst']}%", f"{uasb['ef_coli']}%"],
        "Salida": [f"{uasb['DBO_out']:.1f}", f"{uasb['DQO_out']:.1f}",
                     f"{uasb['SST_out']:.1f}", f"{uasb['Coli_out']:.2e}"],
    })
    st.dataframe(df_ef_uasb, hide_index=True, use_container_width=True)

    # Biogás + energía
    st.subheader("🔥 Producción de biogás (CH₄) y energía")
    energia = energia_biogas(uasb["Q_CH4"])

    bc1, bc2, bc3, bc4 = st.columns(4)
    bc1.metric("DQO → CH₄", f"{uasb['DQO_metano']:,.0f} kg/d")
    bc2.metric("Q CH₄", f"{uasb['Q_CH4']:,.1f} m³/d")
    bc3.metric("CH₄ anual", f"{uasb['Q_CH4']*365:,.0f} m³/año")
    bc4.metric("Energía térmica", f"{energia['energia_termica_d']:,.0f} kWh/d")

    e1, e2, e3 = st.columns(3)
    e1.metric("Energía eléctrica (η=35%)", f"{energia['energia_electrica_d']:,.0f} kWh/d")
    e2.metric("Potencia continua", f"{energia['potencia_kw']:,.1f} kW")
    e3.metric("MWh/año", f"{energia['MWh_anio']:,.1f} MWh")

    st.caption("Eficiencia eléctrica de cogeneración asumida 35%. "
                "PCI CH₄ = 9,97 kWh/m³.")


# ============================================================
# 12. SELECCIÓN DEL PRIMARIO ACTIVO
# ============================================================
if tipo_primario == "Sedimentador 1°":
    primary = sed
elif tipo_primario == "Laguna anaerobia":
    primary = ana
else:
    primary = uasb


# -----------------------------------------------------------------
# TAB FACULTATIVA
# -----------------------------------------------------------------
with tab_fac:
    st.header("🌱 Laguna Facultativa")
    st.caption(f"Recibe efluente de {tipo_primario} (DBO entrada = {primary['DBO_out']:.1f} mg/L)")

    col_a, col_b = st.columns([1, 2])
    with col_a:
        Ls_calc = carga_superficial_facultativa(T_agua)
        st.metric("Ls (Mara)", f"{Ls_calc:.0f} kg/ha·d")
        H_fac = st.slider("H (m)", 1.5, 2.0, 1.8, 0.1, key="h_fac")
        LW_fac = st.slider("L/W", 2.0, 4.0, 2.5, 0.1, key="lw_fac")
        k20_fac = st.number_input("k20 (d⁻¹)", value=0.30,
                                   min_value=0.20, max_value=0.40, step=0.01)
        SS_eff_fac = st.number_input("SS efluente (mg/L)", value=80,
                                       min_value=40, max_value=150, step=5)
        DBOpart_factor = st.number_input("mg DBO / mg SS", value=0.35,
                                          min_value=0.30, max_value=0.40, step=0.01)
        n_fac = st.slider("Lagunas en paralelo", 1, 4, 2, 1, key="n_fac")
        talud_fac = st.slider("Talud m:1", 1.5, 3.0, 2.0, 0.1, key="t_fac")

    DBO_in_fac = primary["DBO_out"]
    DQO_in_fac = primary["DQO_out"]
    SST_in_fac = primary["SST_out"]
    Coli_in_fac = primary["Coli_out"]

    L_org_fac = Q_d * DBO_in_fac / 1000
    A_fac_ha = L_org_fac / Ls_calc
    A_fac = A_fac_ha * 10_000
    A_fac_unit = A_fac / n_fac
    V_fac = A_fac * H_fac
    V_fac_unit = V_fac / n_fac
    TRH_fac = V_fac / Q_d
    L_f, W_f = dimensiones_rectangulares(A_fac_unit, LW_fac)

    k_T_fac = k_corregido(k20_fac, T_agua, theta=1.05)
    DBO_sol_fac = remover(DBO_in_fac, k_T_fac, TRH_fac, regimen, n=1, L_W=LW_fac)
    DBO_part_fac = SS_eff_fac * DBOpart_factor
    DBO_out_fac = DBO_sol_fac + DBO_part_fac
    DQO_out_fac = DQO_in_fac * 0.30
    SST_out_fac = SS_eff_fac
    Coli_out_fac = Coli_in_fac * 0.10

    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("Área total", f"{A_fac_ha:.2f} ha")
        m2.metric("Área/módulo", f"{A_fac_unit:,.0f} m²")
        m3.metric("TRH", f"{TRH_fac:.1f} d")
        m4, m5, m6 = st.columns(3)
        m4.metric("L (módulo)", f"{L_f:.1f} m")
        m5.metric("W (módulo)", f"{W_f:.1f} m")
        m6.metric("k corregido", f"{k_T_fac:.3f} d⁻¹")

        if 15 <= TRH_fac <= 45:
            st.success(f"✓ TRH = {TRH_fac:.1f} d (15-45 d)")
        else:
            st.warning(f"⚠ TRH fuera de rango")

    st.subheader("📐 Plano esquemático")
    fig_fac = dibujar_laguna(L_f, W_f, H_fac, talud=talud_fac,
                                tipo="Facultativa", color_agua="#86c5a8")
    st.pyplot(fig_fac)
    plt.close(fig_fac)

    df_ef_fac = pd.DataFrame({
        "Parámetro": ["DBO soluble", "DBO partic.", "DBO total",
                        "DQO", "SST", "Coliformes"],
        "Entrada": [f"{DBO_in_fac:.1f}", "—", f"{DBO_in_fac:.1f}",
                      f"{DQO_in_fac:.1f}", f"{SST_in_fac:.1f}",
                      f"{Coli_in_fac:.2e}"],
        "Salida": [f"{DBO_sol_fac:.1f}", f"{DBO_part_fac:.1f}",
                     f"{DBO_out_fac:.1f}", f"{DQO_out_fac:.1f}",
                     f"{SST_out_fac:.1f}", f"{Coli_out_fac:.2e}"],
    })
    st.dataframe(df_ef_fac, hide_index=True, use_container_width=True)


# -----------------------------------------------------------------
# TAB MADURACIÓN
# -----------------------------------------------------------------
with tab_mad:
    st.header("💧 Lagunas de Maduración")

    col_a, col_b = st.columns([1, 2])
    with col_a:
        n_lag = st.number_input("N° lagunas en serie", min_value=1,
                                  max_value=10, value=3, step=1)
        TRH_unit = st.slider("TRH/laguna (d)", 1.0, 10.0, 3.0, 0.5)
        H_mad = st.slider("H (m)", 0.7, 1.2, 1.0, 0.1, key="h_mad")
        LW_mad = st.slider("L/W", 4.0, 20.0, 10.0, 0.5, key="lw_mad")
        kb20 = st.number_input("Kb 20°C (d⁻¹)", value=2.6, step=0.1)
        theta_kb = st.number_input("θ Kb", value=1.19, step=0.01)
        talud_mad = st.slider("Talud m:1", 1.5, 3.0, 2.0, 0.1, key="t_mad")

    Kb_T = k_corregido(kb20, T_agua, theta_kb)
    TRH_total = TRH_unit * n_lag
    V_mad_unit = Q_d * TRH_unit
    V_mad_total = V_mad_unit * n_lag
    A_mad_unit = V_mad_unit / H_mad
    A_mad_total = A_mad_unit * n_lag
    L_m, W_m = dimensiones_rectangulares(A_mad_unit, LW_mad)

    Coli_out = remover(Coli_out_fac, Kb_T, TRH_unit, regimen,
                        n=int(n_lag), L_W=LW_mad)

    DBO_out_mad = DBO_out_fac * 0.70
    DQO_out_mad = DQO_out_fac * 0.75
    SST_out_mad = SST_out_fac * 0.80

    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("Kb corregido", f"{Kb_T:.2f} d⁻¹")
        m2.metric("TRH total", f"{TRH_total:.1f} d")
        m3.metric("V total", f"{V_mad_total:,.0f} m³")
        m4, m5, m6 = st.columns(3)
        m4.metric("Área/laguna", f"{A_mad_unit:,.0f} m²")
        m5.metric("L × W", f"{L_m:.0f} × {W_m:.0f} m")
        m6.metric("Coli salida", f"{Coli_out:.2e}")

        if Coli_out < 1000:
            st.success("✓ Cumple norma riego (<1000 NMP/100mL)")
        else:
            st.error("✗ NO cumple. Aumenta n° lagunas o TRH.")

        if n_lag < 3:
            st.warning("⚠ Mínimo 3 lagunas (Clase 09)")
        if TRH_unit < 3:
            st.warning("⚠ TRH/laguna ≥ 3 d")

    st.subheader("📐 Plano esquemático")
    fig_mad = dibujar_lagunas_serie(L_m, W_m, H_mad, n_lag)
    st.pyplot(fig_mad)
    plt.close(fig_mad)


# -----------------------------------------------------------------
# TAB LODOS
# -----------------------------------------------------------------
with tab_lod:
    st.header("🪣 Manejo de lodos")
    st.caption("Volumen anual, frecuencia de evacuación y área de lechos de secado")

    # Datos para cada unidad
    lodos_sed = calcular_lodos("Sedimentador 1°", pob_diseno,
                                  sed["A_unit"], sed["H"])
    lodos_ana = calcular_lodos("Laguna anaerobia", pob_diseno,
                                  ana["A_unit"], ana["H"])
    lodos_uasb = calcular_lodos("Reactor UASB", pob_diseno,
                                   uasb["A_unit"], uasb["H"])
    lodos_fac = calcular_lodos("Laguna facultativa", pob_diseno,
                                  A_fac_unit, H_fac)
    lodos_mad = calcular_lodos("Laguna maduración", pob_diseno,
                                  A_mad_unit, H_mad)

    df_lodos = pd.DataFrame([
        {"Unidad": "Sedimentador 1°", **lodos_sed},
        {"Unidad": "Laguna anaerobia", **lodos_ana},
        {"Unidad": "Reactor UASB", **lodos_uasb},
        {"Unidad": "Laguna facultativa", **lodos_fac},
        {"Unidad": "Laguna maduración", **lodos_mad},
    ])

    df_lodos_disp = df_lodos.copy()
    df_lodos_disp["Tasa (m³/hab·año)"] = df_lodos_disp["tasa"]
    df_lodos_disp["V lodo (m³/año)"] = df_lodos_disp["V_anual"].round(0)
    df_lodos_disp["Acum. (cm/año)"] = df_lodos_disp["altura_anual_cm"].round(1)
    df_lodos_disp["Evac. (años)"] = df_lodos_disp["anos_evacuacion"].round(1)
    df_lodos_disp["Masa SST (kg/año)"] = df_lodos_disp["masa_sst_anual"].round(0)
    df_lodos_disp["A. lecho (m²)"] = df_lodos_disp["A_lecho"].round(0)

    st.dataframe(
        df_lodos_disp[["Unidad", "Tasa (m³/hab·año)", "V lodo (m³/año)",
                        "Acum. (cm/año)", "Evac. (años)",
                        "Masa SST (kg/año)", "A. lecho (m²)"]],
        hide_index=True, use_container_width=True,
    )

    # Resumen del tren seleccionado
    st.subheader(f"📋 Resumen para tren actual ({tipo_primario})")
    if tipo_primario == "Sedimentador 1°":
        primario_lodo = lodos_sed
    elif tipo_primario == "Laguna anaerobia":
        primario_lodo = lodos_ana
    else:
        primario_lodo = lodos_uasb

    V_lodo_total = (primario_lodo["V_anual"] +
                     lodos_fac["V_anual"] +
                     lodos_mad["V_anual"] * n_lag)
    A_lecho_total = (primario_lodo["A_lecho"] +
                      lodos_fac["A_lecho"] +
                      lodos_mad["A_lecho"] * n_lag)

    l1, l2, l3 = st.columns(3)
    l1.metric("V lodo total/año", f"{V_lodo_total:,.0f} m³")
    l2.metric("A lechos secado", f"{A_lecho_total:,.0f} m²")
    l3.metric("Frecuencia evac.",
                f"~{primario_lodo['anos_evacuacion']:.1f} años (primario)")

    # Disposición
    st.subheader("♻️ Alternativas de disposición final")
    st.markdown("""
    Las opciones más comunes para los lodos digeridos generados son:

    - **Reúso agrícola** como mejorador de suelos (siempre que cumpla con la normativa de patógenos y metales pesados de la Resolución 1287 de 2014).
    - **Aplicación en relleno sanitario** como cobertura diaria, previa deshidratación en lechos de secado hasta ~30% de sólidos.
    - **Estabilización adicional con cal** antes de disposición si el contenido de patógenos es alto.
    - **Compostaje** mezclando con residuos vegetales para producir abono orgánico.

    El criterio de evacuación es retirar los lodos cuando ocupen **1/3 de la altura útil** de la laguna, lo que se da generalmente cada 5-10 años en sistemas anaerobios bien operados.
    """)

    # Asunciones
    with st.expander("Asunciones del cálculo"):
        st.markdown("""
        - Tasas per cápita basadas en von Sperling & Chernicharo (2005), Clase 09.
        - Densidad lodo digerido: 1.100 kg/m³.
        - Concentración SST en lodo digerido: 5%.
        - Carga lecho de secado: 150 kg SST/m²·año (Crites & Tchobanoglous).
        - Periodo de digestión: ≥ 6 meses.
        """)


# -----------------------------------------------------------------
# TAB SENSIBILIDAD
# -----------------------------------------------------------------
with tab_sens:
    st.header("📈 Análisis de sensibilidad")
    st.caption("Cómo varía el diseño con la temperatura y la concentración de DBO")

    # Sensibilidad a T
    st.subheader("🌡️ Sensibilidad a la temperatura del agua")
    T_range = np.arange(15.0, 30.5, 0.5)
    df_T = pd.DataFrame({
        "T (°C)": T_range,
        "Lv anaerobia (kg/m³·d)": [carga_volumetrica_anaerobia(t) for t in T_range],
        "Ls facultativa (kg/ha·d)": [carga_superficial_facultativa(t) for t in T_range],
        "Ef. DBO anaerobia (%)": [eficiencia_anaerobia_dbo(t) for t in T_range],
        "TRH UASB (h)": [trh_uasb_por_temperatura(t) for t in T_range],
        "Kb coliformes (d⁻¹)": [k_corregido(2.6, t, 1.19) for t in T_range],
    })

    cs1, cs2 = st.columns(2)
    with cs1:
        st.write("**Lv anaerobia y Ls facultativa**")
        st.line_chart(df_T.set_index("T (°C)")[
            ["Lv anaerobia (kg/m³·d)", "Ls facultativa (kg/ha·d)"]
        ])
    with cs2:
        st.write("**Eficiencia DBO anaerobia y TRH UASB**")
        st.line_chart(df_T.set_index("T (°C)")[
            ["Ef. DBO anaerobia (%)", "TRH UASB (h)"]
        ])

    st.write("**Kb coliformes corregido por T**")
    st.line_chart(df_T.set_index("T (°C)")[["Kb coliformes (d⁻¹)"]])

    with st.expander("📊 Tabla completa"):
        st.dataframe(df_T.style.format({
            "Lv anaerobia (kg/m³·d)": "{:.3f}",
            "Ls facultativa (kg/ha·d)": "{:.0f}",
            "Ef. DBO anaerobia (%)": "{:.1f}",
            "TRH UASB (h)": "{:.1f}",
            "Kb coliformes (d⁻¹)": "{:.3f}",
        }), use_container_width=True)

    # Sensibilidad a DBO_in
    st.subheader("📊 Sensibilidad al DBO de entrada")
    DBO_range = np.arange(150, 451, 25)
    Q_fix = Q_d
    Lv_fix = carga_volumetrica_anaerobia(T_agua)
    Ls_fix = carga_superficial_facultativa(T_agua)

    df_DBO = pd.DataFrame({
        "DBO entrada (mg/L)": DBO_range,
        "V anaerobia (m³)": [Q_fix * d / (1000 * Lv_fix) for d in DBO_range],
        "A facultativa (m²)": [Q_fix * d * (1 - eficiencia_anaerobia_dbo(T_agua)/100) /
                                (Ls_fix) * 10 for d in DBO_range],
        "V UASB (m³)": [Q_fix * trh_uasb_por_temperatura(T_agua)/24
                          for _ in DBO_range],
    })

    st.line_chart(df_DBO.set_index("DBO entrada (mg/L)")[
        ["V anaerobia (m³)", "V UASB (m³)"]
    ])
    st.line_chart(df_DBO.set_index("DBO entrada (mg/L)")[
        ["A facultativa (m²)"]
    ])
    st.caption("Volumen UASB no depende de DBO porque se diseña por TRH, no por carga.")


# -----------------------------------------------------------------
# TAB RESUMEN
# -----------------------------------------------------------------
with tab_res:
    st.header("📊 Resumen Ejecutivo")
    st.markdown(f"### Tren: **{tipo_primario}** → Facultativa → Maduración")

    # Diagrama del tren
    st.subheader("🚂 Diagrama del tren de tratamiento")
    unidades_tren = [
        {"nombre": tipo_primario, "V": primary["V"],
          "TRH": (f"{primary['TRH_h']:.1f} h" if "TRH_h" in primary
                  else f"{primary['TRH_d']:.1f} d"),
          "DBO": primary["DBO_out"], "Coli": primary["Coli_out"]},
        {"nombre": "Facultativa", "V": V_fac,
          "TRH": f"{TRH_fac:.1f} d",
          "DBO": DBO_out_fac, "Coli": Coli_out_fac},
        {"nombre": f"Maduración ×{int(n_lag)}", "V": V_mad_total,
          "TRH": f"{TRH_total:.1f} d",
          "DBO": DBO_out_mad, "Coli": Coli_out},
    ]
    fig_tren = dibujar_tren(unidades_tren)
    st.pyplot(fig_tren)
    plt.close(fig_tren)

    # Caída de contaminantes
    df_caida = pd.DataFrame({
        "Etapa": ["Afluente", f"Efl. {tipo_primario}",
                    "Efl. Facultativa", "Efl. Maduración"],
        "DBO₅ (mg/L)": [DBO_in, primary['DBO_out'], DBO_out_fac, DBO_out_mad],
        "DQO (mg/L)": [DQO_in, primary['DQO_out'], DQO_out_fac, DQO_out_mad],
        "SST (mg/L)": [SST_in, primary['SST_out'], SST_out_fac, SST_out_mad],
        "Coliformes": [Coli_in, primary['Coli_out'], Coli_out_fac, Coli_out],
    })
    st.subheader("Caída de contaminantes")
    st.dataframe(
        df_caida.style.format({
            "DBO₅ (mg/L)": "{:.1f}",
            "DQO (mg/L)": "{:.1f}",
            "SST (mg/L)": "{:.1f}",
            "Coliformes": "{:.2e}",
        }),
        hide_index=True, use_container_width=True,
    )

    col1, col2 = st.columns(2)
    with col1:
        st.bar_chart(df_caida.set_index("Etapa")["DBO₅ (mg/L)"])
        st.caption("DBO₅")
        st.bar_chart(df_caida.set_index("Etapa")["SST (mg/L)"])
        st.caption("SST")
    with col2:
        st.bar_chart(df_caida.set_index("Etapa")["DQO (mg/L)"])
        st.caption("DQO")
        df_log = df_caida[["Etapa", "Coliformes"]].copy()
        df_log["log10"] = np.log10(df_log["Coliformes"].clip(lower=1))
        st.bar_chart(df_log.set_index("Etapa")["log10"])
        st.caption("Coliformes (log)")

    # Eficiencias globales
    st.subheader("Eficiencias globales")
    ef_dbo_g = (1 - DBO_out_mad / DBO_in) * 100
    ef_dqo_g = (1 - DQO_out_mad / DQO_in) * 100
    ef_sst_g = (1 - SST_out_mad / SST_in) * 100
    ef_coli_log = (math.log10(Coli_in / Coli_out) if Coli_out > 0 else float("inf"))

    g1, g2, g3, g4 = st.columns(4)
    g1.metric("DBO₅", f"{ef_dbo_g:.1f}%")
    g2.metric("DQO", f"{ef_dqo_g:.1f}%")
    g3.metric("SST", f"{ef_sst_g:.1f}%")
    g4.metric("Coliformes", f"{ef_coli_log:.1f} log")

    # Cumplimiento Res 0631
    st.subheader("📋 Cumplimiento normativo (Res 0631 de 2015)")
    cumpl = verificar_res0631(DBO_out_mad, DQO_out_mad, SST_out_mad, Coli_out)
    df_cumpl = pd.DataFrame(cumpl)
    st.dataframe(
        df_cumpl.style.format({"Valor": "{:.2e}", "Límite": "{:.2e}"}),
        hide_index=True, use_container_width=True,
    )

    # Exportación a Word
    st.subheader("📥 Exportar memoria de cálculo")
    if DOCX_OK:
        if st.button("📄 Generar memoria en Word (.docx)"):
            with st.spinner("Generando documento..."):
                # Preparar contenidos
                contexto = {
                    "municipio": DATOS_BARBOSA["municipio"],
                    "primario": tipo_primario,
                    "datos": {
                        "Población diseño": f"{pob_diseno:,.0f} hab",
                        "Qmd": f"{Qmd:,.1f} m³/d",
                        "QMD": f"{QMD:,.1f} m³/d",
                        "QMH": f"{QMH:,.1f} m³/d",
                        "DBO₅ entrada": f"{DBO_in:.0f} mg/L",
                        "DQO entrada": f"{DQO_in:.0f} mg/L",
                        "SST entrada": f"{SST_in:.0f} mg/L",
                        "Coliformes entrada": f"{Coli_in:.2e}",
                        "T agua": f"{T_agua:.1f} °C",
                        "Régimen hidráulico": regimen,
                        "Área disponible": f"{area_disp:,.0f} m²",
                    },
                    "fig_tren": fig_to_buf(dibujar_tren(unidades_tren)),
                    "memorias": {},
                    "resumen_df": df_caida.applymap(
                        lambda x: f"{x:.2f}" if isinstance(x, float) else x),
                    "cumplimiento": cumpl,
                }

                # Memoria primario
                if tipo_primario == "Sedimentador 1°":
                    fig_p = dibujar_sedimentador(sed["L"], sed["W"], sed["H"], sed["n_paralelo"])
                    contexto["memorias"]["Sedimentador Primario"] = {
                        "fig": fig_to_buf(fig_p),
                        "texto": (
                            f"Área = {sed['A']:,.0f} m² ({sed['n_paralelo']} módulos)\n"
                            f"V = {sed['V']:,.0f} m³ · TRH = {sed['TRH_h']:.2f} h\n"
                            f"L × W = {sed['L']:.1f} × {sed['W']:.1f} m · H = {sed['H']:.1f} m\n"
                            f"DBO salida = {sed['DBO_out']:.1f} mg/L\n"
                            f"DQO salida = {sed['DQO_out']:.1f} mg/L\n"
                            f"SST salida = {sed['SST_out']:.1f} mg/L"
                        ),
                    }
                elif tipo_primario == "Laguna anaerobia":
                    fig_p = dibujar_laguna(ana["L"], ana["W"], ana["H"],
                                              talud=ana["talud"], tipo="Anaerobia",
                                              color_agua="#9bb8d6")
                    contexto["memorias"]["Laguna Anaerobia"] = {
                        "fig": fig_to_buf(fig_p),
                        "texto": (
                            f"Lv = {ana['Lv']:.3f} kg/m³·d · "
                            f"V = {ana['V']:,.0f} m³\n"
                            f"A = {ana['A']:,.0f} m² ({ana['n_paralelo']} lagunas)\n"
                            f"L × W = {ana['L']:.1f} × {ana['W']:.1f} m · "
                            f"H = {ana['H']:.1f} m · talud {ana['talud']:.0f}:1\n"
                            f"TRH = {ana['TRH_d']:.2f} d\n"
                            f"Eficiencia DBO = {ana['ef_dbo']:.0f}%\n"
                            f"DBO salida = {ana['DBO_out']:.1f} mg/L"
                        ),
                    }
                else:
                    fig_p = dibujar_uasb(uasb["forma"], uasb["L"], uasb["W"],
                                           uasb["H"], uasb["D"])
                    contexto["memorias"]["Reactor UASB"] = {
                        "fig": fig_to_buf(fig_p),
                        "texto": (
                            f"Forma: {uasb['forma']} ({uasb['n_paralelo']} módulos)\n"
                            f"V = {uasb['V']:,.0f} m³ · TRH = {uasb['TRH_h']:.1f} h\n"
                            f"H = {uasb['H']:.1f} m · "
                            + (f"D = {uasb['D']:.1f} m"
                                if uasb['forma'] == "Circular"
                                else f"L × W = {uasb['L']:.1f} × {uasb['W']:.1f} m") + "\n"
                            f"Vs (Qmd) = {uasb['Vs_mh']:.2f} m/h · "
                            f"Vs (QMH) = {uasb['Vs_mh_QMH']:.2f} m/h\n"
                            f"Lo = {uasb['Lo']:.2f} kg DQO/m³·d\n"
                            f"DBO salida = {uasb['DBO_out']:.1f} mg/L\n"
                            f"Producción CH₄ = {uasb['Q_CH4']:,.1f} m³/d "
                            f"({uasb['Q_CH4']*365:,.0f} m³/año)\n"
                            f"Energía eléctrica = {energia_biogas(uasb['Q_CH4'])['MWh_anio']:,.1f} MWh/año"
                        ),
                    }

                # Facultativa
                fig_f = dibujar_laguna(L_f, W_f, H_fac, talud=talud_fac,
                                         tipo="Facultativa", color_agua="#86c5a8")
                contexto["memorias"]["Laguna Facultativa"] = {
                    "fig": fig_to_buf(fig_f),
                    "texto": (
                        f"Ls = {Ls_calc:.0f} kg DBO/ha·d (Mara)\n"
                        f"A = {A_fac:,.0f} m² ({n_fac} lagunas)\n"
                        f"V = {V_fac:,.0f} m³ · TRH = {TRH_fac:.2f} d\n"
                        f"L × W = {L_f:.1f} × {W_f:.1f} m · H = {H_fac:.1f} m\n"
                        f"k corregido = {k_T_fac:.3f} d⁻¹\n"
                        f"DBO salida = {DBO_out_fac:.1f} mg/L"
                    ),
                }

                # Maduración
                fig_m = dibujar_lagunas_serie(L_m, W_m, H_mad, n_lag)
                contexto["memorias"]["Lagunas de Maduración"] = {
                    "fig": fig_to_buf(fig_m),
                    "texto": (
                        f"N° lagunas en serie = {int(n_lag)}\n"
                        f"TRH/laguna = {TRH_unit:.1f} d · TRH total = {TRH_total:.1f} d\n"
                        f"Kb corregido = {Kb_T:.2f} d⁻¹\n"
                        f"V/laguna = {V_mad_unit:,.0f} m³ · V total = {V_mad_total:,.0f} m³\n"
                        f"L × W = {L_m:.1f} × {W_m:.1f} m · H = {H_mad:.1f} m\n"
                        f"Coliformes salida = {Coli_out:.2e} NMP/100mL"
                    ),
                }

                doc_bytes = generar_word_memoria(contexto)
                if doc_bytes:
                    st.download_button(
                        label="⬇️ Descargar memoria.docx",
                        data=doc_bytes,
                        file_name="memoria_PTAR_Barbosa.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    st.success("✓ Memoria generada correctamente.")
    else:
        st.info("Para exportar a Word instala: `pip install python-docx`")


# Footer
st.markdown("---")
st.caption("Desarrollado para *Sistemas de Tratamiento de Aguas Residuales* (UdeA). "
           "Basado en Clases 07-09 + Clase UASB (Molina Pérez), Mara (1997), "
           "von Sperling & Chernicharo (2005), van Haandel & Lettinga (1994), "
           "RAS 2017, Resolución 0631 de 2015.")