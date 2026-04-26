"""
Diseño de PTAR — Caso de estudio: Barbosa, Antioquia
Materia: Sistemas de Tratamiento de Aguas Residuales — UdeA

Funcionalidades:
  • 3 alternativas primarias: Sedimentador, Anaerobia, UASB
  • Selección por área disponible
  • Planos esquemáticos con simbología técnica (norte, escala, cuadro)
  • Layout de planta en el lote
  • Tabla de operación y mantenimiento por unidad
  • Comparador con lodos activados (tren alternativo)
  • Manejo de lodos, costos, Res 0631, sensibilidad
  • Aprovechamiento energético del biogás
  • Exportación a Word

Lógica de ingeniería en diseno.py · Tests en test_diseno.py
"""

import io
import math
import datetime as _dt

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon, Rectangle, FancyBboxPatch, Circle

import numpy as np
import pandas as pd
import streamlit as st

from diseno import (
    DATOS_BARBOSA, RES_0631, REUSO_LIMITES, COSTOS_RANGOS, OM_DATA,
    proyeccion_geometrica, calcular_caudales, temperatura_agua,
    carga_volumetrica_anaerobia, eficiencia_anaerobia_dbo,
    carga_superficial_facultativa, k_corregido,
    s_mezcla_completa, s_flujo_piston, s_flujo_disperso,
    numero_dispersion, remover, dimensiones_rectangulares,
    diametro_circular, trh_uasb_por_temperatura, K_metano,
    calcular_lodos, estimar_costos, energia_biogas, verificar_res0631,
    disenar_sedimentador_primario, disenar_laguna_anaerobia,
    disenar_reactor_uasb, disenar_lodos_activados, recomendar_primario,
)

try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


# ============================================================
# CONFIGURACIÓN
# ============================================================
st.set_page_config(page_title="PTAR Barbosa", layout="wide", page_icon="💧")


# ============================================================
# COLORES Y AYUDANTES PARA PLANOS
# ============================================================
COLOR_AGUA = "#cfe8ff"
COLOR_LODO = "#8b6f47"
COLOR_CONCRETO = "#bbbbbb"
COLOR_BIOGAS = "#fff2b3"
COLOR_BORDE = "#222222"


def _setup_axis(ax, title="", xlim=None, ylim=None):
    ax.set_aspect("equal")
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=10, weight="bold")
    if xlim:
        ax.set_xlim(xlim)
    if ylim:
        ax.set_ylim(ylim)


def _add_north_arrow(ax, x, y, size=1.0):
    """Añade rosa de los vientos / norte a un eje."""
    ax.annotate(
        "N", xy=(x, y + size * 0.7), xytext=(x, y - size * 0.3),
        arrowprops=dict(arrowstyle="-|>,head_width=0.4,head_length=0.6",
                          color="#c0392b", lw=2),
        fontsize=11, color="#c0392b", ha="center", va="bottom", weight="bold",
    )
    circle = Circle((x, y), size * 0.15, fill=True, facecolor="white",
                     edgecolor="#c0392b", linewidth=1.5, zorder=10)
    ax.add_patch(circle)


def _add_scale_bar(ax, x, y, length_m, height=0.3):
    """Añade barra de escala gráfica."""
    half = length_m / 2
    # Negro/blanco alternados
    ax.add_patch(Rectangle((x, y), half, height, facecolor="black", edgecolor="black"))
    ax.add_patch(Rectangle((x + half, y), half, height,
                              facecolor="white", edgecolor="black"))
    # Marcas
    ax.plot([x, x], [y, y + height * 1.5], "k-", lw=1)
    ax.plot([x + half, x + half], [y, y + height * 1.5], "k-", lw=1)
    ax.plot([x + length_m, x + length_m], [y, y + height * 1.5], "k-", lw=1)
    # Etiquetas
    ax.text(x, y + height * 2, "0", fontsize=7, ha="center")
    ax.text(x + half, y + height * 2, f"{half:.0f}", fontsize=7, ha="center")
    ax.text(x + length_m, y + height * 2,
              f"{length_m:.0f} m", fontsize=7, ha="center")
    ax.text(x + length_m / 2, y - height * 1.5,
              "ESCALA", fontsize=7, ha="center", style="italic")


def _add_title_block(ax, x, y, w, h, proyecto, unidad, escala="S/E", autor="J. Amud"):
    """Cuadro de información tipo plano de ingeniería."""
    fecha = _dt.date.today().strftime("%Y-%m-%d")
    ax.add_patch(Rectangle((x, y), w, h, fill=True, facecolor="white",
                              edgecolor="black", linewidth=1.5, zorder=5))
    # Líneas internas (dividir en filas)
    rows = 4
    for i in range(1, rows):
        yi = y + h * i / rows
        ax.plot([x, x + w], [yi, yi], "k-", lw=0.8, zorder=6)
    # Línea vertical interna
    ax.plot([x + w * 0.35] * 2, [y, y + h], "k-", lw=0.8, zorder=6)

    rh = h / rows
    label_size = 6.5
    val_size = 7.5
    # Fila superior: proyecto
    ax.text(x + w * 0.025, y + h - rh * 0.5, "PROYECTO",
              fontsize=label_size, va="center", weight="bold")
    ax.text(x + w * 0.38, y + h - rh * 0.5, proyecto,
              fontsize=val_size, va="center")
    # Unidad
    ax.text(x + w * 0.025, y + h - rh * 1.5, "UNIDAD",
              fontsize=label_size, va="center", weight="bold")
    ax.text(x + w * 0.38, y + h - rh * 1.5, unidad,
              fontsize=val_size, va="center")
    # Escala
    ax.text(x + w * 0.025, y + h - rh * 2.5, "ESCALA",
              fontsize=label_size, va="center", weight="bold")
    ax.text(x + w * 0.38, y + h - rh * 2.5, escala,
              fontsize=val_size, va="center")
    # Fecha y autor
    ax.text(x + w * 0.025, y + h - rh * 3.5, "FECHA / AUTOR",
              fontsize=label_size, va="center", weight="bold")
    ax.text(x + w * 0.38, y + h - rh * 3.5, f"{fecha} / {autor}",
              fontsize=val_size, va="center")


def _add_legend(ax, items, x, y, item_h=0.4):
    """Añade leyenda con cuadros de color y descripción."""
    n = len(items)
    box_w = max(len(it[1]) for it in items) * 0.15 + 0.8
    box_h = item_h * n + 0.3
    ax.add_patch(Rectangle((x, y), box_w, box_h, fill=True,
                              facecolor="white", edgecolor="black", linewidth=1, zorder=5))
    ax.text(x + box_w / 2, y + box_h - item_h * 0.3, "LEYENDA",
              fontsize=7.5, weight="bold", ha="center", zorder=6)
    for i, (color, label) in enumerate(items):
        yi = y + box_h - item_h * (i + 1.2)
        ax.add_patch(Rectangle((x + 0.15, yi - item_h * 0.15),
                                  item_h * 0.6, item_h * 0.5,
                                  facecolor=color, edgecolor="black", zorder=6))
        ax.text(x + 0.15 + item_h * 0.7, yi + item_h * 0.05,
                  label, fontsize=7, va="center", zorder=6)


# ============================================================
# PLANOS ESQUEMÁTICOS — con simbología técnica
# ============================================================
def dibujar_sedimentador(L, W, H, n=1, fig_size=(13, 6)):
    """Plano + corte de sedimentador primario rectangular con simbología."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=fig_size,
                                      gridspec_kw={'width_ratios': [1.2, 1]})
    # ---- PLANTA ----
    ax1.add_patch(Rectangle((0, 0), L, W, facecolor=COLOR_AGUA,
                              edgecolor=COLOR_BORDE, linewidth=1.5))
    ax1.add_patch(Rectangle((-L * 0.04, W * 0.4), L * 0.04, W * 0.2,
                              facecolor=COLOR_CONCRETO, edgecolor="black"))
    ax1.annotate("Q", xy=(0, W / 2), xytext=(-L * 0.13, W / 2),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center", ha="right")
    ax1.plot([L * 0.85] * 2, [0, W], "k-", lw=1.2)
    ax1.text(L * 0.85, W * 1.04, "Bafle", ha="center", fontsize=7)
    ax1.plot([L] * 2, [0, W], "k-", lw=2.5)
    ax1.text(L, W * 1.04, "Vertedero", ha="center", fontsize=7)
    ax1.annotate("Q", xy=(L * 1.13, W / 2), xytext=(L, W / 2),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center")
    ax1.add_patch(Rectangle((L * 0.05, W * 0.05), L * 0.78, W * 0.1,
                              facecolor=COLOR_LODO, alpha=0.4))
    ax1.text(L / 2, W * 0.1, "Tolva de lodos", ha="center", va="center", fontsize=7)
    # Dim
    ax1.annotate("", xy=(L, -W * 0.18), xytext=(0, -W * 0.18),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax1.text(L / 2, -W * 0.27, f"L = {L:.1f} m", ha="center",
              fontsize=9, weight="bold")
    ax1.annotate("", xy=(-L * 0.08, W), xytext=(-L * 0.08, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax1.text(-L * 0.16, W / 2, f"W = {W:.1f} m", va="center",
              rotation=90, fontsize=9, weight="bold")
    # Norte
    _add_north_arrow(ax1, L * 1.1, W * 0.95, size=W * 0.12)
    # Escala
    sb_len = max(round(L / 5), 5)
    _add_scale_bar(ax1, 0, -W * 0.42, sb_len, height=W * 0.025)
    title_pl = f"PLANTA — {n} módulo(s) en paralelo" if n > 1 else "PLANTA"
    _setup_axis(ax1, title_pl,
                  xlim=(-L * 0.27, L * 1.27), ylim=(-W * 0.55, W * 1.25))

    # ---- CORTE ----
    pts_tank = [[0, 0], [0, H], [L, H], [L, H * 0.05], [L * 0.5, 0]]
    ax2.add_patch(Polygon(pts_tank, facecolor=COLOR_AGUA,
                            edgecolor=COLOR_BORDE, linewidth=1.5))
    ax2.plot([0, L], [H * 0.92] * 2, "b--", lw=1, alpha=0.6)
    ax2.text(L * 0.97, H * 0.97, "▽", fontsize=10, color="blue")
    ax2.text(0, H * 1.05, "Cota 0.00 = solera",
              fontsize=7, color="gray", style="italic")
    pts_lodo = [[L * 0.05, 0], [L * 0.55, 0], [L * 0.95, H * 0.13], [L * 0.05, H * 0.13]]
    ax2.add_patch(Polygon(pts_lodo, facecolor=COLOR_LODO, alpha=0.7))
    ax2.text(L * 0.4, H * 0.07, "Lodos", ha="center", va="center",
              fontsize=8, color="white")
    ax2.plot([L * 0.85] * 2, [H * 0.4, H * 0.95], "k-", lw=2)
    ax2.annotate("Q", xy=(0, H * 0.6), xytext=(-L * 0.1, H * 0.6),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center")
    ax2.annotate("Q", xy=(L * 1.1, H * 0.85), xytext=(L, H * 0.85),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center")
    ax2.annotate("Lodos", xy=(L * 0.5, -H * 0.18), xytext=(L * 0.5, 0),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="brown"),
                  fontsize=8, color="brown", ha="center")
    ax2.annotate("", xy=(-L * 0.05, H), xytext=(-L * 0.05, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax2.text(-L * 0.13, H / 2, f"H = {H:.1f} m", va="center",
              rotation=90, fontsize=9, weight="bold")
    ax2.annotate("", xy=(L, -H * 0.4), xytext=(0, -H * 0.4),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax2.text(L / 2, -H * 0.5, f"L = {L:.1f} m", ha="center",
              fontsize=9, weight="bold")
    _setup_axis(ax2, "CORTE LONGITUDINAL",
                  xlim=(-L * 0.22, L * 1.22), ylim=(-H * 0.7, H * 1.25))

    plt.suptitle("SEDIMENTADOR PRIMARIO", fontsize=12, weight="bold", y=1.02)
    plt.tight_layout()
    return fig


def dibujar_laguna(L, W, H, talud=2, tipo="Anaerobia",
                     color_agua=COLOR_AGUA, fig_size=(13, 6)):
    """Plano + corte de laguna trapezoidal con simbología técnica."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=fig_size,
                                      gridspec_kw={'width_ratios': [1.2, 1]})
    L_bot = max(L - 2 * H * talud, L * 0.1)
    W_bot = max(W - 2 * H * talud, W * 0.1)

    # ---- PLANTA ----
    ax1.add_patch(Rectangle((0, 0), L, W, facecolor=color_agua,
                              edgecolor=COLOR_BORDE, linewidth=1.5))
    ox, oy = (L - L_bot) / 2, (W - W_bot) / 2
    ax1.add_patch(Rectangle((ox, oy), L_bot, W_bot,
                              fill=False, edgecolor="gray",
                              linewidth=1, linestyle="--"))
    ax1.text(L / 2, W * 0.5, "Espejo de agua", ha="center", va="center",
              fontsize=8, color="#555", style="italic")
    ax1.plot(0, W * 0.5, "o", color="blue", markersize=10)
    ax1.annotate("Q", xy=(0, W * 0.5), xytext=(-L * 0.1, W * 0.5),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center")
    ax1.plot(L, W * 0.5, "s", color="red", markersize=8)
    ax1.annotate("Q", xy=(L * 1.1, W * 0.5), xytext=(L, W * 0.5),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center")
    # Dim
    ax1.annotate("", xy=(L, -W * 0.18), xytext=(0, -W * 0.18),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax1.text(L / 2, -W * 0.27, f"L (corona) = {L:.1f} m",
              ha="center", fontsize=9, weight="bold")
    ax1.annotate("", xy=(-L * 0.08, W), xytext=(-L * 0.08, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax1.text(-L * 0.16, W / 2, f"W = {W:.1f} m",
              va="center", rotation=90, fontsize=9, weight="bold")
    # Norte y escala
    _add_north_arrow(ax1, L * 1.1, W * 0.95, size=W * 0.12)
    sb_len = max(round(L / 5), 5)
    _add_scale_bar(ax1, 0, -W * 0.42, sb_len, height=W * 0.025)
    _setup_axis(ax1, "PLANTA",
                  xlim=(-L * 0.27, L * 1.27), ylim=(-W * 0.55, W * 1.25))

    # ---- CORTE TRANSVERSAL ----
    pts = [[0, 0], [(W - W_bot) / 2, H], [W - (W - W_bot) / 2, H], [W, 0]]
    ax2.add_patch(Polygon(pts, facecolor=color_agua,
                            edgecolor=COLOR_BORDE, linewidth=1.5))
    bx_l = (W - W_bot) / 2 * 1.05
    bx_r = W - (W - W_bot) / 2 * 1.05
    ax2.plot([bx_l, bx_r], [H * 0.93] * 2, "b--", lw=1, alpha=0.6)
    ax2.text(W * 0.05, H * 1.05, "▽ NA", fontsize=8, color="blue")
    sludge_h = H * 0.12
    inset = sludge_h * talud
    pts_lodo = [
        [(W - W_bot) / 2, 0],
        [W - (W - W_bot) / 2, 0],
        [W - (W - W_bot) / 2 - inset, sludge_h],
        [(W - W_bot) / 2 + inset, sludge_h],
    ]
    ax2.add_patch(Polygon(pts_lodo, facecolor=COLOR_LODO, alpha=0.7))
    ax2.text(W / 2, sludge_h * 0.4, "Lodos digeridos",
              ha="center", va="center", fontsize=8, color="white")
    ax2.text((W - W_bot) / 4, H * 0.5, f"1\n―\n{talud:.0f}",
              ha="center", va="center", fontsize=8,
              bbox=dict(boxstyle="round", facecolor="white", alpha=0.7))
    # Suelo natural
    ax2.fill_between([-W * 0.05, W * 1.05], -H * 0.05, 0,
                       color="#deb887", alpha=0.4, hatch="//")
    ax2.plot([-W * 0.05, W * 1.05], [0] * 2, "k-", lw=0.8)
    # Dim
    ax2.annotate("", xy=(-W * 0.08, H), xytext=(-W * 0.08, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax2.text(-W * 0.15, H / 2, f"H = {H:.1f} m",
              va="center", rotation=90, fontsize=9, weight="bold")
    ax2.annotate("", xy=(W, -H * 0.3), xytext=(0, -H * 0.3),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax2.text(W / 2, -H * 0.42, f"W (corona) = {W:.1f} m",
              ha="center", fontsize=9, weight="bold")
    ax2.annotate("Q", xy=(0, H * 0.7), xytext=(-W * 0.1, H * 0.7),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center")
    ax2.annotate("Q", xy=(W * 1.1, H * 0.85), xytext=(W, H * 0.85),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=10, color="blue", va="center")
    _setup_axis(ax2, f"CORTE TRANSVERSAL — talud {talud:.0f}:1",
                  xlim=(-W * 0.22, W * 1.22), ylim=(-H * 0.65, H * 1.25))

    plt.suptitle(f"LAGUNA {tipo.upper()}", fontsize=12, weight="bold", y=1.02)
    plt.tight_layout()
    return fig


def dibujar_uasb(forma, L, W, H, D, fig_size=(8, 7)):
    """Corte vertical de UASB con simbología."""
    fig, ax = plt.subplots(figsize=fig_size)
    if forma == "Circular":
        a = D
        labels = (f"D = {D:.1f} m", f"H = {H:.1f} m")
        title = f"REACTOR UASB CIRCULAR"
    else:
        a = W
        labels = (f"W = {W:.1f} m", f"H = {H:.1f} m")
        title = f"REACTOR UASB RECTANGULAR"

    ax.add_patch(Rectangle((0, 0), a, H, facecolor=COLOR_AGUA,
                              edgecolor=COLOR_BORDE, linewidth=1.5))
    ax.add_patch(Rectangle((0, 0), a, H * 0.20,
                              facecolor="#5d3a1f", alpha=0.85))
    ax.text(a / 2, H * 0.10, "Lecho de lodo (40-100 g/L)",
              ha="center", va="center", fontsize=9, color="white", weight="bold")
    ax.add_patch(Rectangle((0, H * 0.20), a, H * 0.30,
                              facecolor=COLOR_LODO, alpha=0.55))
    ax.text(a / 2, H * 0.35, "Manto de lodo (10-30 g/L)",
              ha="center", va="center", fontsize=9, color="black")
    pts_sep = [[a * 0.05, H * 0.65], [a * 0.30, H * 0.85],
                [a * 0.70, H * 0.85], [a * 0.95, H * 0.65]]
    ax.add_patch(Polygon(pts_sep, facecolor=COLOR_CONCRETO,
                          edgecolor="black", linewidth=1.2))
    ax.text(a / 2, H * 0.78, "Separador G-L-S",
              ha="center", va="center", fontsize=9, weight="bold")
    ax.add_patch(Rectangle((a * 0.30, H * 0.85), a * 0.40, H * 0.13,
                              facecolor=COLOR_BIOGAS, edgecolor="black", linewidth=1))
    ax.text(a / 2, H * 0.92, "Biogás", ha="center", va="center",
              fontsize=9, weight="bold")
    ax.annotate("CH₄ + CO₂", xy=(a * 0.5, H * 1.15), xytext=(a * 0.5, H * 0.98),
                  arrowprops=dict(arrowstyle="->", lw=2, color="orange"),
                  fontsize=9, color="orange", ha="center", va="bottom", weight="bold")
    for i in range(1, 5):
        x = a * i / 5
        ax.annotate("", xy=(x, H * 0.05), xytext=(x, -H * 0.1),
                      arrowprops=dict(arrowstyle="->", color="blue", lw=1))
    ax.text(a / 2, -H * 0.15, "Influente (distribuidor)",
              ha="center", fontsize=9, color="blue")
    ax.annotate("Efluente", xy=(a * 1.18, H * 0.78),
                  xytext=(a, H * 0.78),
                  arrowprops=dict(arrowstyle="->", lw=2, color="green"),
                  fontsize=9, color="green", va="center", weight="bold")
    ax.annotate("", xy=(-a * 0.06, H), xytext=(-a * 0.06, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax.text(-a * 0.13, H / 2, labels[1],
              va="center", rotation=90, fontsize=10, weight="bold")
    ax.annotate("", xy=(a, -H * 0.3), xytext=(0, -H * 0.3),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax.text(a / 2, -H * 0.4, labels[0], ha="center", fontsize=10, weight="bold")
    # Escala
    sb_len = max(round(a / 4), 2)
    _add_scale_bar(ax, 0, -H * 0.55, sb_len, height=H * 0.025)
    # Cota
    ax.text(0, -H * 0.05, "Cota 0.00 = solera",
              fontsize=7, color="gray", style="italic")
    _setup_axis(ax, title,
                  xlim=(-a * 0.27, a * 1.35), ylim=(-H * 0.7, H * 1.3))
    plt.tight_layout()
    return fig


def dibujar_lagunas_serie(L, W, H, n, fig_size=(14, 5)):
    """Vista en planta de n lagunas en serie con simbología."""
    fig, ax = plt.subplots(figsize=fig_size)
    sep = L * 0.05
    for i in range(int(n)):
        x = i * (L + sep)
        ax.add_patch(Rectangle((x, 0), L, W, facecolor=COLOR_AGUA,
                                  edgecolor=COLOR_BORDE, linewidth=1.5))
        ax.text(x + L / 2, W / 2, f"Mad. {i + 1}",
                  ha="center", va="center", fontsize=10, weight="bold")
        if i < n - 1:
            ax.annotate("", xy=(x + L + sep, W / 2), xytext=(x + L, W / 2),
                          arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"))
        else:
            ax.annotate("Q", xy=(x + L + sep, W / 2), xytext=(x + L, W / 2),
                          arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                          fontsize=11, color="blue", va="center")
    ax.annotate("Q", xy=(0, W / 2), xytext=(-sep, W / 2),
                  arrowprops=dict(arrowstyle="->", lw=1.5, color="blue"),
                  fontsize=11, color="blue", va="center", ha="right")
    ax.annotate("", xy=(L, -W * 0.15), xytext=(0, -W * 0.15),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax.text(L / 2, -W * 0.25, f"L = {L:.1f} m",
              ha="center", fontsize=9)
    ax.annotate("", xy=(-sep, W), xytext=(-sep, 0),
                  arrowprops=dict(arrowstyle="<->", color="black"))
    ax.text(-sep * 1.5, W / 2, f"W = {W:.1f} m",
              va="center", rotation=90, fontsize=9)
    total_L = n * L + (n - 1) * sep
    # Norte
    _add_north_arrow(ax, total_L * 1.02, W * 0.95, size=W * 0.15)
    # Escala
    sb_len = max(round(L / 4), 5)
    _add_scale_bar(ax, 0, -W * 0.5, sb_len, height=W * 0.03)
    _setup_axis(ax,
                  f"LAGUNAS DE MADURACIÓN — {int(n)} en serie · H = {H:.1f} m",
                  xlim=(-sep * 2, total_L + sep * 3),
                  ylim=(-W * 0.65, W * 1.25))
    plt.tight_layout()
    return fig


def dibujar_lodos_activados(la_dict, fig_size=(13, 5)):
    """Esquema del tren de lodos activados: reactor + sed sec + desinfección."""
    fig, ax = plt.subplots(figsize=fig_size)
    L_aer, W_aer, H_aer = la_dict["L_aer"], la_dict["W_aer"], la_dict["H_aer"]
    L_sec, W_sec = la_dict["L_sec"], la_dict["W_sec"]
    sep = max(L_aer, L_sec) * 0.15

    # Reactor de aireación
    ax.add_patch(Rectangle((0, 0), L_aer, W_aer, facecolor="#a8d8ea",
                              edgecolor=COLOR_BORDE, linewidth=1.5))
    ax.text(L_aer / 2, W_aer / 2, f"Reactor\nAireación\n{la_dict['V_aer']:,.0f} m³",
              ha="center", va="center", fontsize=10, weight="bold")
    # Aireadores (círculos)
    for i in range(4):
        cx = L_aer * (i + 0.5) / 4
        cy = W_aer * 0.85
        ax.add_patch(Circle((cx, cy), W_aer * 0.05, facecolor="orange",
                              edgecolor="black", zorder=5))
    ax.text(L_aer / 2, -W_aer * 0.2, "Aireadores", fontsize=8,
              ha="center", color="orange")

    # Sedimentador secundario
    sx = L_aer + sep
    ax.add_patch(Rectangle((sx, 0), L_sec, W_sec, facecolor="#86c5a8",
                              edgecolor=COLOR_BORDE, linewidth=1.5))
    ax.text(sx + L_sec / 2, W_sec / 2,
              f"Sed. 2°\n{la_dict['A_sec']:.0f} m²",
              ha="center", va="center", fontsize=10, weight="bold")

    # Cloración
    cx = sx + L_sec + sep
    cw = L_sec * 0.5
    ax.add_patch(Rectangle((cx, 0), cw, W_sec * 0.7, facecolor="#fcbad3",
                              edgecolor=COLOR_BORDE, linewidth=1.5))
    ax.text(cx + cw / 2, W_sec * 0.35, "Cloración",
              ha="center", va="center", fontsize=9, weight="bold")

    # Flechas
    ax.annotate("Q", xy=(0, W_aer / 2), xytext=(-sep, W_aer / 2),
                  arrowprops=dict(arrowstyle="->", lw=2, color="blue"),
                  fontsize=11, color="blue", va="center", ha="right")
    ax.annotate("", xy=(sx, W_aer / 2), xytext=(L_aer, W_aer / 2),
                  arrowprops=dict(arrowstyle="->", lw=2, color="blue"))
    ax.annotate("", xy=(cx, W_sec / 2), xytext=(sx + L_sec, W_sec / 2),
                  arrowprops=dict(arrowstyle="->", lw=2, color="blue"))
    ax.annotate("Efluente", xy=(cx + cw + sep, W_sec / 2),
                  xytext=(cx + cw, W_sec / 2),
                  arrowprops=dict(arrowstyle="->", lw=2, color="green"),
                  fontsize=10, color="green", va="center")
    # Recirculación
    ax.annotate("", xy=(L_aer / 2, -W_aer * 0.4),
                  xytext=(sx + L_sec * 0.3, -W_aer * 0.4),
                  arrowprops=dict(arrowstyle="->", lw=1.2, color="brown",
                                    connectionstyle="arc3,rad=-0.3"))
    ax.text((L_aer + sx) / 2, -W_aer * 0.55,
              "Recirculación de lodos (Qr)",
              ha="center", fontsize=8, color="brown")

    # Norte y escala
    L_total = cx + cw
    _add_north_arrow(ax, L_total * 1.05, W_aer, size=W_aer * 0.15)
    sb_len = max(round(L_aer / 3), 2)
    _add_scale_bar(ax, 0, -W_aer * 0.85, sb_len, height=W_aer * 0.04)
    _setup_axis(ax, "TREN DE LODOS ACTIVADOS",
                  xlim=(-sep * 2, L_total + sep * 2),
                  ylim=(-W_aer * 1.1, W_aer * 1.3))
    plt.tight_layout()
    return fig


def dibujar_tren(unidades, fig_size=(13, 4.5)):
    """Diagrama esquemático del tren de tratamiento."""
    fig, ax = plt.subplots(figsize=fig_size)
    n = len(unidades)
    box_w, box_h = 1.0, 0.7
    sep = 0.6
    y = 0
    colores = ["#a8d8ea", "#aa96da", "#fcbad3", "#ffffd2", "#a0e8af"]
    for i, u in enumerate(unidades):
        x = i * (box_w + sep)
        c = colores[i % len(colores)]
        box = FancyBboxPatch((x, y), box_w, box_h,
                              boxstyle="round,pad=0.02",
                              facecolor=c, edgecolor="black", linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + box_w / 2, y + box_h * 0.7, u["nombre"],
                  ha="center", fontsize=10, weight="bold")
        ax.text(x + box_w / 2, y + box_h * 0.45,
                  f"V = {u['V']:,.0f} m³", ha="center", fontsize=8)
        ax.text(x + box_w / 2, y + box_h * 0.25,
                  f"TRH = {u['TRH']}", ha="center", fontsize=8)
        ax.text(x + box_w / 2, y - 0.15,
                  f"DBO = {u['DBO']:.0f} mg/L",
                  ha="center", fontsize=8, color="darkred")
        ax.text(x + box_w / 2, y - 0.30,
                  f"Coli = {u['Coli']:.1e}",
                  ha="center", fontsize=7, color="darkblue")
        if i < n - 1:
            ax.annotate("", xy=(x + box_w + sep, y + box_h / 2),
                          xytext=(x + box_w, y + box_h / 2),
                          arrowprops=dict(arrowstyle="->", lw=2, color="blue"))
    ax.annotate("Afluente",
                  xy=(0, y + box_h / 2), xytext=(-sep * 0.6, y + box_h / 2),
                  arrowprops=dict(arrowstyle="->", lw=2, color="blue"),
                  fontsize=9, color="blue", va="center", ha="right")
    ax.annotate("Efluente",
                  xy=((n - 1) * (box_w + sep) + box_w + sep * 0.6, y + box_h / 2),
                  xytext=((n - 1) * (box_w + sep) + box_w, y + box_h / 2),
                  arrowprops=dict(arrowstyle="->", lw=2, color="green"),
                  fontsize=9, color="green", va="center")
    _setup_axis(ax, "TREN DE TRATAMIENTO",
                  xlim=(-sep * 1.5, n * (box_w + sep) + sep),
                  ylim=(-0.6, box_h + 0.3))
    plt.tight_layout()
    return fig


def dibujar_layout_planta(area_disp, primario, fac_data, mad_data, n_mad,
                            lechos_area, fig_size=(14, 10)):
    """Layout en planta de la PTAR completa en el lote disponible."""
    # Lote 1.5:1
    W_lot = math.sqrt(area_disp / 1.5)
    L_lot = 1.5 * W_lot

    fig, ax = plt.subplots(figsize=fig_size)
    # Lote
    ax.add_patch(Rectangle((0, 0), L_lot, W_lot, fill=True,
                              facecolor="#f5e6d3", alpha=0.3,
                              edgecolor="black", linewidth=2.5, linestyle="--"))
    ax.text(L_lot / 2, W_lot + L_lot * 0.025,
              f"LOTE — {area_disp:,.0f} m² ({L_lot:.0f} × {W_lot:.0f} m)",
              ha="center", fontsize=11, weight="bold")

    margin = max(L_lot, W_lot) * 0.02
    units_drawn = []  # tracking what fits

    # Caseta administrativa
    cw, ch = 12, 8
    cx, cy = margin, W_lot - margin - ch
    ax.add_patch(Rectangle((cx, cy), cw, ch,
                              facecolor="#e0e0e0", edgecolor="black"))
    ax.text(cx + cw / 2, cy + ch / 2, "Caseta\nadmón.",
              ha="center", va="center", fontsize=8)
    cx += cw + margin

    # Pretratamiento
    pw, ph = 18, 8
    ax.add_patch(Rectangle((cx, cy), pw, ph,
                              facecolor="#fff2b3", edgecolor="black"))
    ax.text(cx + pw / 2, cy + ph / 2, "Pretrata-\nmiento",
              ha="center", va="center", fontsize=8)
    cx_after_pre = cx + pw + margin

    # Primario(s)
    primL, primW = primario["L"], primario["W"]
    n_p = primario.get("n_paralelo", 1)
    py_top = W_lot - margin
    px = cx_after_pre
    for i in range(int(n_p)):
        py = py_top - primW
        # Si se sale del lote a la derecha, salto a otra fila
        if px + primL > L_lot - margin:
            px = cx_after_pre
            py_top = py - margin
            py = py_top - primW
        ax.add_patch(Rectangle((px, py), primL, primW,
                                  facecolor="#a8d8ea", edgecolor="black"))
        ax.text(px + primL / 2, py + primW / 2,
                  f"{primario['nombre']}\n{i + 1}\n{primL:.0f}×{primW:.0f}m",
                  ha="center", va="center", fontsize=8, weight="bold")
        units_drawn.append((px, py, primL, primW, primario["nombre"] + str(i + 1)))
        px += primL + margin
    primario_bottom = min(py for (_, py, _, _, _) in units_drawn[-int(n_p):])

    # Facultativa(s)
    facL, facW = fac_data["L"], fac_data["W"]
    n_f = fac_data.get("n_paralelo", 1)
    fy_top = primario_bottom - margin
    fx = margin
    for i in range(int(n_f)):
        fy = fy_top - facW
        if fx + facL > L_lot - margin:
            fx = margin
            fy_top = fy - margin
            fy = fy_top - facW
        ax.add_patch(Rectangle((fx, fy), facL, facW,
                                  facecolor="#86c5a8", edgecolor="black"))
        ax.text(fx + facL / 2, fy + facW / 2,
                  f"Facultativa {i + 1}\n{facL:.0f}×{facW:.0f}m",
                  ha="center", va="center", fontsize=9, weight="bold")
        units_drawn.append((fx, fy, facL, facW, f"Fac{i+1}"))
        fx += facL + margin
    fac_bottom = min(units_drawn[-int(n_f) + j][1] for j in range(int(n_f)))

    # Maduración
    madL, madW = mad_data["L"], mad_data["W"]
    my_top = fac_bottom - margin
    mx = margin
    for i in range(int(n_mad)):
        my = my_top - madW
        if mx + madL > L_lot - margin:
            mx = margin
            my_top = my - margin
            my = my_top - madW
        ax.add_patch(Rectangle((mx, my), madL, madW,
                                  facecolor="#fcbad3", edgecolor="black"))
        ax.text(mx + madL / 2, my + madW / 2,
                  f"Mad. {i + 1}",
                  ha="center", va="center", fontsize=9, weight="bold")
        units_drawn.append((mx, my, madL, madW, f"Mad{i+1}"))
        mx += madL + margin

    # Lechos secado (esquina inferior derecha)
    lecho_w = math.sqrt(lechos_area * 1.5)
    lecho_h = math.sqrt(lechos_area / 1.5)
    lx = L_lot - lecho_w - margin
    ly = margin
    ax.add_patch(Rectangle((lx, ly), lecho_w, lecho_h,
                              facecolor="#d4a574", edgecolor="black"))
    ax.text(lx + lecho_w / 2, ly + lecho_h / 2,
              f"Lechos\nsecado\n{lechos_area:.0f} m²",
              ha="center", va="center", fontsize=8)

    # Verificar overflow
    overflow_count = 0
    for (px_, py_, pL_, pW_, name) in units_drawn:
        if px_ < 0 or py_ < 0 or (px_ + pL_) > L_lot or (py_ + pW_) > W_lot:
            overflow_count += 1
            ax.add_patch(Rectangle((px_, py_), pL_, pW_,
                                      fill=False, edgecolor="red",
                                      linewidth=2, linestyle=":", zorder=20))

    # Norte
    _add_north_arrow(ax, L_lot * 1.06, W_lot * 0.93, size=W_lot * 0.07)
    # Escala
    sb_len = max(round(L_lot / 6 / 10) * 10, 10)
    _add_scale_bar(ax, margin, -L_lot * 0.04, sb_len, height=W_lot * 0.012)
    # Cuadro de información
    tb_w, tb_h = L_lot * 0.32, W_lot * 0.16
    tb_x, tb_y = L_lot - tb_w - margin, -L_lot * 0.18
    _add_title_block(ax, tb_x, tb_y, tb_w, tb_h,
                       proyecto="PTAR BARBOSA, ANTIOQUIA",
                       unidad="LAYOUT GENERAL", escala="VARIABLE",
                       autor="J. Amud — UdeA")

    # Leyenda
    leg_x = -L_lot * 0.16
    leg_y = W_lot * 0.05
    items = [
        ("#e0e0e0", "Caseta admón."),
        ("#fff2b3", "Pretratamiento"),
        ("#a8d8ea", primario["nombre"]),
        ("#86c5a8", "Facultativa"),
        ("#fcbad3", "Maduración"),
        ("#d4a574", "Lechos secado"),
    ]
    leg_box_h = len(items) * 1.5 + 2
    leg_box_w = 8
    ax.add_patch(Rectangle((leg_x - 1, leg_y - 1), leg_box_w + 2,
                              leg_box_h + 1, fill=True,
                              facecolor="white", edgecolor="black", linewidth=1))
    ax.text(leg_x + leg_box_w / 2, leg_y + leg_box_h - 0.5,
              "LEYENDA", fontsize=8, weight="bold", ha="center")
    for i, (color, label) in enumerate(items):
        yi = leg_y + leg_box_h - 2 - i * 1.5
        ax.add_patch(Rectangle((leg_x, yi), 1.2, 1,
                                  facecolor=color, edgecolor="black"))
        ax.text(leg_x + 1.5, yi + 0.5, label, fontsize=7, va="center")

    # Status overflow
    if overflow_count > 0:
        ax.text(L_lot / 2, W_lot * 1.06,
                  f"⚠ {overflow_count} unidad(es) exceden el lote disponible",
                  ha="center", fontsize=10, color="red", weight="bold")

    margin_view = max(L_lot, W_lot) * 0.20
    ax.set_xlim(-margin_view, L_lot + margin_view)
    ax.set_ylim(-margin_view * 1.2, W_lot + margin_view * 0.5)
    ax.set_aspect("equal")
    ax.axis("off")
    plt.tight_layout()
    return fig


# ============================================================
# EXPORTACIÓN A WORD
# ============================================================
def fig_to_buf(fig, dpi=120):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    return buf


def generar_word_memoria(contexto):
    if not DOCX_OK:
        return None
    doc = Document()
    doc.add_heading("Memoria de Cálculo - PTAR Barbosa", 0)
    p = doc.add_paragraph()
    p.add_run("Materia: ").bold = True
    p.add_run("Sistemas de Tratamiento de Aguas Residuales · UdeA\n")
    p.add_run("Caso: ").bold = True
    p.add_run(f"{contexto['municipio']}\n")
    p.add_run("Tren seleccionado: ").bold = True
    p.add_run(f"{contexto['primario']} → Facultativa → Maduración")

    doc.add_heading("1. Datos del proyecto", 1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "Parámetro"
    t.rows[0].cells[1].text = "Valor"
    for k, v in contexto["datos"].items():
        row = t.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    doc.add_heading("2. Tren de tratamiento", 1)
    if contexto.get("fig_tren") is not None:
        doc.add_picture(contexto["fig_tren"], width=Inches(6.5))

    doc.add_heading("3. Memorias por unidad", 1)
    for nombre, mem in contexto["memorias"].items():
        doc.add_heading(nombre, 2)
        if mem.get("fig") is not None:
            doc.add_picture(mem["fig"], width=Inches(6.0))
        for linea in mem["texto"].split("\n"):
            if linea.strip():
                doc.add_paragraph(linea)

    doc.add_heading("4. Resumen", 1)
    if "resumen_df" in contexto:
        df = contexto["resumen_df"]
        t2 = doc.add_table(rows=1, cols=len(df.columns))
        t2.style = "Light Grid Accent 1"
        for i, c in enumerate(df.columns):
            t2.rows[0].cells[i].text = str(c)
        for _, row in df.iterrows():
            cells = t2.add_row().cells
            for i, c in enumerate(df.columns):
                cells[i].text = str(row[c])

    doc.add_heading("5. Cumplimiento Res 0631 de 2015", 1)
    if "cumplimiento" in contexto:
        df = pd.DataFrame(contexto["cumplimiento"])
        t3 = doc.add_table(rows=1, cols=len(df.columns))
        t3.style = "Light Grid Accent 1"
        for i, c in enumerate(df.columns):
            t3.rows[0].cells[i].text = str(c)
        for _, row in df.iterrows():
            cells = t3.add_row().cells
            for i, c in enumerate(df.columns):
                cells[i].text = str(row[c])

    doc.add_paragraph()
    doc.add_paragraph(
        "Generado automáticamente. Mara (1997), von Sperling & Chernicharo (2005), "
        "van Haandel & Lettinga (1994), RAS 2017, Res 0631/2015."
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# SIDEBAR
# ============================================================
st.sidebar.title("⚙️ Parámetros de diseño")
st.sidebar.caption(f"📍 Caso: {DATOS_BARBOSA['municipio']}")

with st.sidebar.expander("📥 Datos del afluente", expanded=False):
    DBO_in = st.number_input("DBO₅ (mg/L)", value=DATOS_BARBOSA["DBO5"],
                              min_value=1.0, step=10.0)
    DQO_in = st.number_input("DQO (mg/L)", value=DATOS_BARBOSA["DQO"],
                              min_value=1.0, step=10.0)
    SST_in = st.number_input("SST (mg/L)", value=DATOS_BARBOSA["SST"],
                              min_value=1.0, step=10.0)
    Coli_in = st.number_input("Coliformes fecales (NMP/100mL)",
                               value=DATOS_BARBOSA["coliformes_fecales"],
                               format="%.2e", min_value=1.0)

with st.sidebar.expander("🌡️ Condiciones ambientales", expanded=True):
    T_aire = st.slider("T aire mes más frío (°C)",
                       min_value=10.0, max_value=30.0,
                       value=DATOS_BARBOSA["T_aire_mes_frio"], step=0.5)
    T_agua_yanez = temperatura_agua(T_aire)
    T_agua = st.slider("T agua del agua residual (°C)",
                        min_value=10.0, max_value=30.0,
                        value=round(T_agua_yanez, 1), step=0.5,
                        help=f"Yáñez: T_agua = {T_agua_yanez:.1f} °C")

with st.sidebar.expander("🌊 Régimen hidráulico", expanded=False):
    regimen = st.selectbox("Modelo hidráulico",
                            ["Mezcla completa", "Flujo pistón", "Flujo disperso"])

with st.sidebar.expander("🏗️ Tratamiento primario", expanded=True):
    area_disp = st.number_input("📐 Área disponible (m²)",
                                  value=DATOS_BARBOSA["area_disponible"],
                                  min_value=0.0, step=500.0)
    tipo_primario = st.radio("Primario activo",
                              ["Sedimentador 1°", "Laguna anaerobia",
                               "Reactor UASB"], index=2)

with st.sidebar.expander("👥 Demografía", expanded=False):
    pob_actual = st.number_input("Población actual (hab)",
                                   value=DATOS_BARBOSA["poblacion_actual"],
                                   min_value=1, step=500)
    tasa = st.number_input("Tasa crecimiento (%/año)",
                             value=DATOS_BARBOSA["tasa_crecimiento"],
                             min_value=0.0, step=0.1)
    horizonte = st.number_input("Horizonte (años)",
                                  value=DATOS_BARBOSA["horizonte_diseno"],
                                  min_value=1, step=1)
    dotacion = st.number_input("Dotación (L/hab·d)",
                                 value=DATOS_BARBOSA["dotacion_neta"],
                                 min_value=1, step=10)
    factor_r = st.number_input("Factor de retorno",
                                 value=DATOS_BARBOSA["factor_retorno"],
                                 min_value=0.05, max_value=1.0, step=0.05)
    k1_in = st.number_input("k1", value=DATOS_BARBOSA["k1"],
                              min_value=1.0, step=0.05)
    k2_in = st.number_input("k2", value=DATOS_BARBOSA["k2"],
                              min_value=1.0, step=0.05)


# ============================================================
# VALIDACIONES
# ============================================================
errores = []
if DBO_in <= 0 or DQO_in <= 0 or SST_in <= 0:
    errores.append("DBO, DQO y SST deben ser positivas.")
if Coli_in <= 0:
    errores.append("Coliformes deben ser positivos.")
if DBO_in > DQO_in:
    errores.append("DBO no puede ser mayor que DQO.")
if pob_actual <= 0 or dotacion <= 0:
    errores.append("Población y dotación deben ser positivas.")
if not (0 < factor_r <= 1):
    errores.append("Factor de retorno entre 0 y 1.")
if k1_in < 1 or k2_in < 1:
    errores.append("k1 y k2 ≥ 1.")
if T_agua <= 0:
    errores.append("Temperatura del agua > 0.")
if area_disp < 0:
    errores.append("Área disponible ≥ 0.")
if errores:
    st.error("❌ Errores en datos de entrada:")
    for er in errores:
        st.error(er)
    st.stop()


# ============================================================
# CÁLCULOS PRELIMINARES
# ============================================================
pob_diseno = proyeccion_geometrica(pob_actual, tasa, horizonte)
Qmd, QMD, QMH = calcular_caudales(pob_diseno, dotacion, factor_r, k1_in, k2_in)
Q_d = Qmd

sed_pre = disenar_sedimentador_primario(Qmd, QMH, DBO_in, DQO_in, SST_in, Coli_in)
ana_pre = disenar_laguna_anaerobia(Q_d, DBO_in, DQO_in, SST_in, Coli_in, T_agua)
uasb_pre = disenar_reactor_uasb(Q_d, QMH, DBO_in, DQO_in, SST_in, Coli_in, T_agua)


# ============================================================
# ENCABEZADO
# ============================================================
st.title("💧 Diseño de PTAR — Caso Barbosa, Antioquia")
st.markdown(f"### Tren: **{tipo_primario}** → Facultativa → Maduración")
c1, c2, c3, c4, c5 = st.columns(5)
c1.metric("Población", f"{pob_diseno:,.0f} hab")
c2.metric("Qmd", f"{Qmd:,.1f} m³/d")
c3.metric("QMD", f"{QMD:,.1f} m³/d")
c4.metric("QMH", f"{QMH:,.1f} m³/d")
c5.metric("Área disp.", f"{area_disp:,.0f} m²")


# ============================================================
# PESTAÑAS
# ============================================================
tabs = st.tabs([
    "📍 Caudales", "📐 Selección + Costos",
    "🏛️ Sedimentador", "🌊 Anaerobia", "⚗️ UASB",
    "🌱 Facultativa", "💧 Maduración",
    "🪣 Lodos", "🗺️ Layout", "⚙️ Lodos Activados",
    "🛠️ O&M", "📈 Sensibilidad", "📊 Resumen",
])
(tab0, tab_sel, tab_sed, tab_ana, tab_uasb, tab_fac, tab_mad,
 tab_lod, tab_lay, tab_la, tab_om, tab_sens, tab_res) = tabs


# -----------------------------------------------------------------
# TAB 0 — Caudales
# -----------------------------------------------------------------
with tab0:
    st.header("Proyección poblacional y caudales")
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Proyección geométrica")
        st.latex(r"P_f = P_0 \cdot (1 + r/100)^n")
        st.dataframe(pd.DataFrame({
            "Año": [DATOS_BARBOSA["anio_actual"],
                     DATOS_BARBOSA["anio_actual"] + horizonte],
            "Población": [int(pob_actual), int(round(pob_diseno))],
        }), hide_index=True, use_container_width=True)
    with col2:
        st.subheader("Caudales (RAS 2017)")
        st.dataframe(pd.DataFrame({
            "Caudal": ["Qmd", "QMD", "QMH"],
            "m³/d": [f"{Qmd:,.1f}", f"{QMD:,.1f}", f"{QMH:,.1f}"],
            "L/s": [f"{Qmd*1000/86400:,.1f}",
                    f"{QMD*1000/86400:,.1f}",
                    f"{QMH*1000/86400:,.1f}"],
        }), hide_index=True, use_container_width=True)
    st.info("ℹ️ Lagunas y UASB con Qmd · Sedimentador con QMH (RAS Art. 189).")


# -----------------------------------------------------------------
# TAB SELECCIÓN
# -----------------------------------------------------------------
with tab_sel:
    st.header("📐 Selección de tratamiento primario")
    rec, status, just, viables = recomendar_primario(
        area_disp, sed_pre, ana_pre, uasb_pre)

    df_comp = pd.DataFrame({
        "Alternativa": ["Sedimentador 1°", "Laguna anaerobia", "Reactor UASB"],
        "Área req (m²)": [sed_pre["A"], ana_pre["A"], uasb_pre["A"]],
        "Volumen (m³)": [sed_pre["V"], ana_pre["V"], uasb_pre["V"]],
        "TRH": [f"{sed_pre['TRH_h']:.1f} h",
                 f"{ana_pre['TRH_d']:.1f} d",
                 f"{uasb_pre['TRH_h']:.1f} h"],
        "Ef. DBO (%)": [sed_pre["ef_dbo"], ana_pre["ef_dbo"], uasb_pre["ef_dbo"]],
        "¿Cabe?": [
            "✅" if sed_pre["A"] <= area_disp else "❌",
            "✅" if ana_pre["A"] <= area_disp else "❌",
            "✅" if uasb_pre["A"] <= area_disp else "❌",
        ],
    })
    st.dataframe(df_comp.style.format({"Área req (m²)": "{:,.0f}",
                                         "Volumen (m³)": "{:,.0f}"}),
                  hide_index=True, use_container_width=True)

    if status == "danger":
        st.error(f"⚠️ {rec}")
    else:
        st.success(f"🎯 **Recomendación:** {rec}")
        st.info(just)
    if rec != tipo_primario and status != "danger":
        st.warning(f"💡 Has seleccionado {tipo_primario} pero la recomendación es {rec}")

    st.subheader("💰 Costos (US$/hab — Clase 09)")
    cs = estimar_costos("Sedimentador 1°", pob_diseno)
    ca = estimar_costos("Laguna anaerobia", pob_diseno)
    cu = estimar_costos("Reactor UASB", pob_diseno)
    cf = estimar_costos("Laguna facultativa", pob_diseno)
    cm = estimar_costos("Lagunas maduración", pob_diseno)
    cla = estimar_costos("Lodos activados", pob_diseno)

    df_costos = pd.DataFrame({
        "Unidad": ["Sedimentador 1°", "Laguna anaerobia", "Reactor UASB",
                    "Lodos activados", "Laguna facultativa", "Lagunas maduración"],
        "Constr. mín": [cs["construccion_min"], ca["construccion_min"],
                          cu["construccion_min"], cla["construccion_min"],
                          cf["construccion_min"], cm["construccion_min"]],
        "Constr. máx": [cs["construccion_max"], ca["construccion_max"],
                          cu["construccion_max"], cla["construccion_max"],
                          cf["construccion_max"], cm["construccion_max"]],
        "O&M anual": [cs["om_anual_med"], ca["om_anual_med"],
                        cu["om_anual_med"], cla["om_anual_med"],
                        cf["om_anual_med"], cm["om_anual_med"]],
        "Total 25 años": [cs["total_25_med"], ca["total_25_med"],
                             cu["total_25_med"], cla["total_25_med"],
                             cf["total_25_med"], cm["total_25_med"]],
    })
    st.dataframe(df_costos.style.format({c: "{:,.0f}" for c in df_costos.columns
                                           if c != "Unidad"}),
                  hide_index=True, use_container_width=True)


# -----------------------------------------------------------------
# TAB SEDIMENTADOR
# -----------------------------------------------------------------
with tab_sed:
    st.header("🏛️ Sedimentador Primario")
    col_a, col_b = st.columns([1, 2])
    with col_a:
        TDS_med = st.slider("TDS Qmd (m³/m²·d)", 30, 50, 40, 1)
        TDS_pico = st.slider("TDS QMH (m³/m²·d)", 70, 130, 100, 5)
        H_sed = st.slider("Profundidad H (m)", 2.5, 4.0, 3.0, 0.1, key="h_sed")
        LW_sed = st.slider("Relación L/W", 1.5, 15.0, 2.5, 0.5, key="lw_sed")
        n_sed = st.slider("Unidades en paralelo", 1, 4, 2, 1, key="n_sed")
    sed = disenar_sedimentador_primario(
        Qmd, QMH, DBO_in, DQO_in, SST_in, Coli_in,
        TDS_med=TDS_med, TDS_pico=TDS_pico, H=H_sed, LW=LW_sed, n_paralelo=n_sed)
    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("Área total", f"{sed['A']:,.0f} m²")
        m2.metric("Área/módulo", f"{sed['A_unit']:,.0f} m²")
        m3.metric("TRH", f"{sed['TRH_h']:.2f} h")
        m4, m5, m6 = st.columns(3)
        m4.metric("L (módulo)", f"{sed['L']:.1f} m")
        m5.metric("W (módulo)", f"{sed['W']:.1f} m")
        m6.metric("V total", f"{sed['V']:,.0f} m³")
        if 1.5 <= sed["TRH_h"] <= 2.5:
            st.success(f"✓ TRH = {sed['TRH_h']:.2f} h (RAS 1.5-2.5 h)")
        else:
            st.warning(f"⚠ TRH fuera de rango")
        if sed["A"] <= area_disp:
            st.success(f"✓ Cabe ({sed['A']/area_disp*100:.1f}%)")
        else:
            st.error(f"✗ NO cabe")
    st.subheader("📐 Plano esquemático")
    fig_sed = dibujar_sedimentador(sed["L"], sed["W"], sed["H"], sed["n_paralelo"])
    st.pyplot(fig_sed)
    plt.close(fig_sed)


# -----------------------------------------------------------------
# TAB ANAEROBIA
# -----------------------------------------------------------------
with tab_ana:
    st.header("🌊 Laguna Anaerobia")
    col_a, col_b = st.columns([1, 2])
    with col_a:
        Lv_calc = carga_volumetrica_anaerobia(T_agua)
        st.metric("Lv (Mara)", f"{Lv_calc:.3f} kg/m³·d")
        Lv_ana = st.number_input("Lv adoptado", value=round(Lv_calc, 3),
                                   min_value=0.05, max_value=0.50, step=0.01)
        H_ana = st.slider("H (m)", 3.0, 5.0, 4.0, 0.1, key="h_ana")
        LW_ana = st.slider("L/W", 1.0, 3.0, 2.0, 0.1, key="lw_ana")
        talud_ana = st.slider("Talud", 1.5, 3.0, 2.0, 0.1, key="t_ana")
        n_ana = st.slider("Lagunas en paralelo", 1, 4, 2, 1, key="n_ana")
    ana = disenar_laguna_anaerobia(
        Q_d, DBO_in, DQO_in, SST_in, Coli_in, T_agua,
        Lv=Lv_ana, H=H_ana, LW=LW_ana, talud=talud_ana, n_paralelo=n_ana)
    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("V total", f"{ana['V']:,.0f} m³")
        m2.metric("Área total", f"{ana['A']:,.0f} m²")
        m3.metric("TRH", f"{ana['TRH_d']:.2f} d")
        m4, m5, m6 = st.columns(3)
        m4.metric("L (módulo)", f"{ana['L']:.1f} m")
        m5.metric("W (módulo)", f"{ana['W']:.1f} m")
        m6.metric("Carga DBO", f"{ana['L_org']:,.0f} kg/d")
        if 3 <= ana["TRH_d"] <= 6:
            st.success(f"✓ TRH = {ana['TRH_d']:.2f} d (3-6 d)")
        else:
            st.warning(f"⚠ TRH fuera de rango")
        if ana["A"] <= area_disp:
            st.success(f"✓ Cabe ({ana['A']/area_disp*100:.1f}%)")
        else:
            st.error(f"✗ NO cabe")
    st.subheader("📐 Plano esquemático")
    fig_ana = dibujar_laguna(ana["L"], ana["W"], ana["H"],
                                talud=ana["talud"], tipo="Anaerobia",
                                color_agua="#9bb8d6")
    st.pyplot(fig_ana)
    plt.close(fig_ana)


# -----------------------------------------------------------------
# TAB UASB
# -----------------------------------------------------------------
with tab_uasb:
    st.header("⚗️ Reactor UASB")
    col_a, col_b = st.columns([1, 2])
    with col_a:
        forma_uasb = st.radio("Forma", ["Circular", "Rectangular"], index=0)
        TRH_calc = trh_uasb_por_temperatura(T_agua)
        st.metric("TRH recomendado", f"{TRH_calc:.1f} h")
        TRH_uasb = st.slider("TRH (h)", 4.0, 14.0, TRH_calc, 0.5, key="trh_u")
        H_uasb = st.slider("H (m)", 4.0, 6.0, 5.0, 0.1, key="h_u")
        LW_uasb = st.slider("L/W (rectangular)", 1.0, 3.0, 1.0, 0.1, key="lw_u")
        n_uasb = st.slider("Reactores en paralelo", 1, 6, 2, 1, key="n_u")
        Y_acid = st.number_input("Y acidogénica", value=0.15, step=0.01)
        Y_metano = st.number_input("Y metanogénica", value=0.03, step=0.01)
    uasb = disenar_reactor_uasb(
        Q_d, QMH, DBO_in, DQO_in, SST_in, Coli_in, T_agua,
        TRH_h=TRH_uasb, H=H_uasb, LW=LW_uasb,
        forma=forma_uasb, n_paralelo=n_uasb,
        Y_acid=Y_acid, Y_metano=Y_metano)
    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("V total", f"{uasb['V']:,.0f} m³")
        m2.metric("V/módulo", f"{uasb['V_unit']:,.0f} m³")
        m3.metric("TRH", f"{uasb['TRH_h']:.1f} h")
        m4, m5, m6 = st.columns(3)
        m4.metric("Vs (Qmd)", f"{uasb['Vs_mh']:.2f} m/h")
        m5.metric("Vs (QMH)", f"{uasb['Vs_mh_QMH']:.2f} m/h")
        m6.metric("Lo", f"{uasb['Lo']:.2f} kg/m³·d")
        if uasb["TRH_h"] >= 4:
            st.success(f"✓ TRH ≥ 4 h")
        else:
            st.error(f"✗ TRH < 4 h")
        if 0.5 <= uasb["Vs_mh"] <= 0.7:
            st.success(f"✓ Vs(Qmd) en rango (0.5-0.7)")
        else:
            st.warning(f"⚠ Vs(Qmd) fuera de rango")
        if uasb["Vs_mh_QMH"] <= 1.1:
            st.success(f"✓ Vs(QMH) ≤ 1.1 m/h")
        else:
            st.error(f"✗ Vs(QMH) > 1.1: riesgo lavado")
        if uasb["A"] <= area_disp:
            st.success(f"✓ Cabe ({uasb['A']/area_disp*100:.1f}%)")
        else:
            st.error(f"✗ NO cabe")
    st.subheader("📐 Plano esquemático")
    fig_uasb = dibujar_uasb(forma_uasb, uasb["L"], uasb["W"],
                                uasb["H"], uasb["D"])
    st.pyplot(fig_uasb)
    plt.close(fig_uasb)
    st.subheader("🔥 Producción de biogás y energía")
    energia = energia_biogas(uasb["Q_CH4"])
    bc1, bc2, bc3, bc4 = st.columns(4)
    bc1.metric("DQO → CH₄", f"{uasb['DQO_metano']:,.0f} kg/d")
    bc2.metric("CH₄/día", f"{uasb['Q_CH4']:,.1f} m³/d")
    bc3.metric("Energía elec.", f"{energia['energia_electrica_d']:,.0f} kWh/d")
    bc4.metric("MWh/año", f"{energia['MWh_anio']:,.1f}")


# ============================================================
# SELECCIÓN DEL PRIMARIO
# ============================================================
if tipo_primario == "Sedimentador 1°":
    primary = sed
elif tipo_primario == "Laguna anaerobia":
    primary = ana
else:
    primary = uasb


# -----------------------------------------------------------------
# TAB FACULTATIVA
# -----------------------------------------------------------------
with tab_fac:
    st.header("🌱 Laguna Facultativa")
    st.caption(f"Recibe efluente de {tipo_primario} (DBO={primary['DBO_out']:.1f} mg/L)")
    col_a, col_b = st.columns([1, 2])
    with col_a:
        Ls_calc = carga_superficial_facultativa(T_agua)
        st.metric("Ls (Mara)", f"{Ls_calc:.0f} kg/ha·d")
        H_fac = st.slider("H (m)", 1.5, 2.0, 1.8, 0.1, key="h_fac")
        LW_fac = st.slider("L/W", 2.0, 4.0, 2.5, 0.1, key="lw_fac")
        k20_fac = st.number_input("k20 (d⁻¹)", value=0.30,
                                   min_value=0.20, max_value=0.40, step=0.01)
        SS_eff_fac = st.number_input("SS efluente", value=80,
                                       min_value=40, max_value=150, step=5)
        DBOpart_factor = st.number_input("mg DBO/mg SS", value=0.35,
                                          min_value=0.30, max_value=0.40, step=0.01)
        n_fac = st.slider("Lagunas en paralelo", 1, 4, 2, 1, key="n_fac")
        talud_fac = st.slider("Talud", 1.5, 3.0, 2.0, 0.1, key="t_fac")
    DBO_in_fac = primary["DBO_out"]
    DQO_in_fac = primary["DQO_out"]
    SST_in_fac = primary["SST_out"]
    Coli_in_fac = primary["Coli_out"]
    L_org_fac = Q_d * DBO_in_fac / 1000
    A_fac_ha = L_org_fac / Ls_calc
    A_fac = A_fac_ha * 10_000
    A_fac_unit = A_fac / n_fac
    V_fac = A_fac * H_fac
    V_fac_unit = V_fac / n_fac
    TRH_fac = V_fac / Q_d
    L_f, W_f = dimensiones_rectangulares(A_fac_unit, LW_fac)
    k_T_fac = k_corregido(k20_fac, T_agua, theta=1.05)
    DBO_sol_fac = remover(DBO_in_fac, k_T_fac, TRH_fac, regimen, n=1, L_W=LW_fac)
    DBO_part_fac = SS_eff_fac * DBOpart_factor
    DBO_out_fac = DBO_sol_fac + DBO_part_fac
    DQO_out_fac = DQO_in_fac * 0.30
    SST_out_fac = SS_eff_fac
    Coli_out_fac = Coli_in_fac * 0.10
    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("Área", f"{A_fac_ha:.2f} ha")
        m2.metric("V", f"{V_fac:,.0f} m³")
        m3.metric("TRH", f"{TRH_fac:.1f} d")
        m4, m5, m6 = st.columns(3)
        m4.metric("L módulo", f"{L_f:.1f} m")
        m5.metric("W módulo", f"{W_f:.1f} m")
        m6.metric("k corregido", f"{k_T_fac:.3f} d⁻¹")
        if 15 <= TRH_fac <= 45:
            st.success(f"✓ TRH = {TRH_fac:.1f} d en rango")
        else:
            st.warning(f"⚠ TRH fuera de rango (15-45 d)")
    st.subheader("📐 Plano esquemático")
    fig_fac = dibujar_laguna(L_f, W_f, H_fac, talud=talud_fac,
                                tipo="Facultativa", color_agua="#86c5a8")
    st.pyplot(fig_fac)
    plt.close(fig_fac)


# -----------------------------------------------------------------
# TAB MADURACIÓN
# -----------------------------------------------------------------
with tab_mad:
    st.header("💧 Lagunas de Maduración")
    col_a, col_b = st.columns([1, 2])
    with col_a:
        n_lag = st.number_input("N° en serie", min_value=1, max_value=10,
                                  value=3, step=1)
        TRH_unit = st.slider("TRH/laguna (d)", 1.0, 10.0, 3.0, 0.5)
        H_mad = st.slider("H (m)", 0.7, 1.2, 1.0, 0.1, key="h_m")
        LW_mad = st.slider("L/W", 4.0, 20.0, 10.0, 0.5, key="lw_m")
        kb20 = st.number_input("Kb 20°C (d⁻¹)", value=2.6, step=0.1)
        theta_kb = st.number_input("θ Kb", value=1.19, step=0.01)
    Kb_T = k_corregido(kb20, T_agua, theta_kb)
    TRH_total = TRH_unit * n_lag
    V_mad_unit = Q_d * TRH_unit
    V_mad_total = V_mad_unit * n_lag
    A_mad_unit = V_mad_unit / H_mad
    A_mad_total = A_mad_unit * n_lag
    L_m, W_m = dimensiones_rectangulares(A_mad_unit, LW_mad)
    Coli_out = remover(Coli_out_fac, Kb_T, TRH_unit, regimen,
                        n=int(n_lag), L_W=LW_mad)
    DBO_out_mad = DBO_out_fac * 0.70
    DQO_out_mad = DQO_out_fac * 0.75
    SST_out_mad = SST_out_fac * 0.80
    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("Kb corregido", f"{Kb_T:.2f} d⁻¹")
        m2.metric("TRH total", f"{TRH_total:.1f} d")
        m3.metric("V total", f"{V_mad_total:,.0f} m³")
        m4, m5, m6 = st.columns(3)
        m4.metric("Área/laguna", f"{A_mad_unit:,.0f} m²")
        m5.metric("L × W", f"{L_m:.0f} × {W_m:.0f}")
        m6.metric("Coli salida", f"{Coli_out:.2e}")
        if Coli_out < 1000:
            st.success("✓ Cumple norma riego")
        else:
            st.error("✗ NO cumple")
    st.subheader("📐 Plano esquemático")
    fig_mad = dibujar_lagunas_serie(L_m, W_m, H_mad, n_lag)
    st.pyplot(fig_mad)
    plt.close(fig_mad)


# -----------------------------------------------------------------
# TAB LODOS
# -----------------------------------------------------------------
with tab_lod:
    st.header("🪣 Manejo de lodos")
    lodos_sed = calcular_lodos("Sedimentador 1°", pob_diseno,
                                  sed["A_unit"], sed["H"])
    lodos_ana = calcular_lodos("Laguna anaerobia", pob_diseno,
                                  ana["A_unit"], ana["H"])
    lodos_uasb = calcular_lodos("Reactor UASB", pob_diseno,
                                   uasb["A_unit"], uasb["H"])
    lodos_fac = calcular_lodos("Laguna facultativa", pob_diseno,
                                  A_fac_unit, H_fac)
    lodos_mad = calcular_lodos("Laguna maduración", pob_diseno,
                                  A_mad_unit, H_mad)
    df_lodos = pd.DataFrame([
        {"Unidad": "Sedimentador 1°", **lodos_sed},
        {"Unidad": "Laguna anaerobia", **lodos_ana},
        {"Unidad": "Reactor UASB", **lodos_uasb},
        {"Unidad": "Laguna facultativa", **lodos_fac},
        {"Unidad": "Laguna maduración", **lodos_mad},
    ])
    df_lodos["V lodo (m³/año)"] = df_lodos["V_anual"].round(0)
    df_lodos["Acum. (cm/año)"] = df_lodos["altura_anual_cm"].round(1)
    df_lodos["Evac. (años)"] = df_lodos["anos_evacuacion"].round(1)
    df_lodos["Masa SST (kg/año)"] = df_lodos["masa_sst_anual"].round(0)
    df_lodos["A. lecho (m²)"] = df_lodos["A_lecho"].round(0)
    st.dataframe(
        df_lodos[["Unidad", "tasa", "V lodo (m³/año)", "Acum. (cm/año)",
                    "Evac. (años)", "Masa SST (kg/año)", "A. lecho (m²)"]],
        hide_index=True, use_container_width=True)
    st.markdown("""
    **Disposición final**: Reúso agrícola (Res 1287/2014), aplicación en relleno
    sanitario (cobertura), estabilización con cal, o compostaje con residuos vegetales.
    """)


# -----------------------------------------------------------------
# TAB LAYOUT (NEW)
# -----------------------------------------------------------------
with tab_lay:
    st.header("🗺️ Layout de la planta en el lote")
    st.caption("Plano de implantación de todas las unidades en el lote disponible")

    # Compute total
    A_total = primary["A"] + A_fac + A_mad_total
    A_lecho_total = (lodos_sed["A_lecho"] if tipo_primario == "Sedimentador 1°"
                     else lodos_ana["A_lecho"] if tipo_primario == "Laguna anaerobia"
                     else lodos_uasb["A_lecho"]) + lodos_fac["A_lecho"]

    c1, c2, c3 = st.columns(3)
    c1.metric("Área total tren", f"{A_total:,.0f} m²")
    c2.metric("Área disponible", f"{area_disp:,.0f} m²")
    c3.metric("Lechos secado", f"{A_lecho_total:,.0f} m²")
    if A_total + A_lecho_total > area_disp:
        st.error(f"⚠ El tren completo ({A_total + A_lecho_total:,.0f} m²) "
                  f"excede el lote disponible ({area_disp:,.0f} m²). "
                  "Considera reducir TRH de maduración, número de lagunas o "
                  "aumentar el área del lote.")
    else:
        ocupacion = (A_total + A_lecho_total) / area_disp * 100
        st.success(f"✓ El tren cabe ({ocupacion:.1f}% del lote ocupado)")

    primario_layout = {
        "L": primary["L"], "W": primary["W"],
        "n_paralelo": primary.get("n_paralelo", 1),
        "nombre": tipo_primario.replace(" 1°", "").replace("Reactor ", ""),
    }
    fac_layout = {"L": L_f, "W": W_f, "n_paralelo": n_fac}
    mad_layout = {"L": L_m, "W": W_m}
    fig_lay = dibujar_layout_planta(area_disp, primario_layout, fac_layout,
                                       mad_layout, n_lag, A_lecho_total)
    st.pyplot(fig_lay)
    plt.close(fig_lay)

    st.info("ℹ️ Las unidades en línea roja punteada indican que exceden los "
            "límites del lote. El layout es esquemático; la disposición real "
            "debe considerar accesos vehiculares (mín. 6 m de ancho), "
            "distancia entre unidades (mín. 5 m), zona de mantenimiento y "
            "topografía del terreno.")


# -----------------------------------------------------------------
# TAB LODOS ACTIVADOS (NEW - tren alternativo)
# -----------------------------------------------------------------
with tab_la:
    st.header("⚙️ Comparador: Lagunas vs Lodos Activados")
    st.caption("Tren alternativo: " + tipo_primario + " → Lodos Activados → Cloración")

    col_a, col_b = st.columns([1, 2])
    with col_a:
        TRH_la = st.slider("TRH aireación (h)", 4.0, 12.0, 6.0, 0.5)
        H_la = st.slider("H reactor (m)", 3.0, 6.0, 4.5, 0.1, key="h_la")
        SSV_la = st.number_input("SSV reactor (mg/L)", value=3000,
                                   min_value=1500, max_value=5000, step=100)
        theta_c_la = st.number_input("Edad lodos θc (d)", value=10,
                                       min_value=5, max_value=20)
        TDS_sec = st.number_input("TDS sed sec (m³/m²·d)", value=20,
                                    min_value=15, max_value=30)
        Y_obs = st.number_input("Y_obs (kg SSV/kg DBO)", value=0.4,
                                  min_value=0.2, max_value=0.6, step=0.05)

    la = disenar_lodos_activados(
        Q_d, primary["DBO_out"], primary["DQO_out"],
        primary["SST_out"], primary["Coli_out"],
        TRH_h=TRH_la, H_aer=H_la, SSV=SSV_la,
        theta_c=theta_c_la, Y_obs=Y_obs, TDS_sec=TDS_sec,
    )

    with col_b:
        m1, m2, m3 = st.columns(3)
        m1.metric("V reactor", f"{la['V_aer']:,.0f} m³")
        m2.metric("A reactor", f"{la['A_aer']:,.0f} m²")
        m3.metric("F/M", f"{la['F_M']:.2f}")
        m4, m5, m6 = st.columns(3)
        m4.metric("L × W reactor", f"{la['L_aer']:.0f} × {la['W_aer']:.0f} m")
        m5.metric("A sed sec", f"{la['A_sec']:,.0f} m²")
        m6.metric("A total", f"{la['A_total']:,.0f} m²")
        m7, m8, m9 = st.columns(3)
        m7.metric("Lodo producido", f"{la['lodo_kg_d']:,.0f} kg/d")
        m8.metric("OR", f"{la['OR_kg_d']:,.0f} kg O₂/d")
        m9.metric("Energía", f"{la['energia_kwh_d']:,.0f} kWh/d")

        if 0.2 <= la["F_M"] <= 0.5:
            st.success(f"✓ F/M = {la['F_M']:.2f} en rango (0.2-0.5)")
        else:
            st.warning(f"⚠ F/M fuera de rango")

    st.subheader("📐 Plano esquemático del tren")
    fig_la = dibujar_lodos_activados(la)
    st.pyplot(fig_la)
    plt.close(fig_la)

    # Comparación lado a lado
    st.subheader("📊 Comparación: Lagunas vs Lodos Activados")
    A_lagunas = A_fac + A_mad_total
    V_lagunas = V_fac + V_mad_total
    DBO_out_LA = la["DBO_out"]
    Coli_out_LA = la["Coli_out"] * 0.01  # asumiendo cloración remueve 2 log

    df_comp_tren = pd.DataFrame({
        "Métrica": ["Área secundario+terciario (m²)",
                      "Volumen total (m³)",
                      "DBO efluente (mg/L)",
                      "SST efluente (mg/L)",
                      "Coliformes efluente",
                      "Energía (kWh/d)",
                      "Lodo (kg/d)",
                      "Costo construcción (US$/hab)",
                      "Costo O&M (US$/hab·año)",
                      "Tiempo arranque",
                      "Operación"],
        "Lagunas (Fac + Mad)": [
            f"{A_lagunas:,.0f}",
            f"{V_lagunas:,.0f}",
            f"{DBO_out_mad:.1f}",
            f"{SST_out_mad:.1f}",
            f"{Coli_out:.2e}",
            "0 (pasivo)",
            f"~{(lodos_fac['masa_sst_anual'] + lodos_mad['masa_sst_anual']*n_lag)/365:,.0f}",
            "23-45 (rango medio)",
            "1.5-3.0",
            "Inmediato",
            "Simple, baja"],
        "Lodos Activados": [
            f"{la['A_total']:,.0f}",
            f"{la['V_aer'] + la['V_sec']:,.0f}",
            f"{DBO_out_LA:.1f}",
            f"{la['SST_out']:.1f}",
            f"{Coli_out_LA:.2e}",
            f"{la['energia_kwh_d']:,.0f}",
            f"{la['lodo_kg_d']:,.0f}",
            "40-80",
            "3.0-6.0",
            "2-4 semanas",
            "Compleja, alta"],
    })
    st.dataframe(df_comp_tren, hide_index=True, use_container_width=True)

    st.markdown("""
    **Análisis comparativo:**

    Las **lagunas** ofrecen una huella mayor pero costos de O&M muy bajos, son robustas a
    sobrecargas, no requieren energía y están adaptadas al clima cálido de Barbosa.
    Producción de lodo baja porque los procesos son extensivos (lentos) y los lodos quedan
    digeridos en la propia laguna.

    Los **lodos activados** ocupan ~10× menos área y producen un efluente de mayor calidad,
    pero requieren operadores capacitados, energía continua para aireación (típicamente
    0.5-1.0 kWh/m³), generan grandes volúmenes de lodo activo que requieren disposición
    posterior y son sensibles a cargas tóxicas o variaciones bruscas.

    Para un municipio como Barbosa con suficiente terreno y clima favorable, las lagunas
    son típicamente la opción **costo-efectiva**. Lodos activados es preferible en zonas
    urbanas con poco espacio o donde se requiere un efluente de alta calidad para reúso
    industrial o riego no restringido sin lagunas adicionales.
    """)


# -----------------------------------------------------------------
# TAB O&M (NEW)
# -----------------------------------------------------------------
with tab_om:
    st.header("🛠️ Operación y Mantenimiento")
    st.caption("Parámetros de monitoreo, frecuencias y acciones por unidad")

    unidad_om = st.selectbox("Unidad",
                              list(OM_DATA.keys()),
                              index=2)  # UASB por defecto

    df_om = pd.DataFrame(
        OM_DATA[unidad_om],
        columns=["Parámetro", "Frecuencia", "Rango/valor de referencia",
                  "Acción si fuera de rango"],
    )
    st.dataframe(df_om, hide_index=True, use_container_width=True)

    # Tabla resumida del tren completo
    st.subheader("📋 Cronograma O&M del tren actual")
    unidades_tren_om = []
    if tipo_primario == "Sedimentador 1°":
        unidades_tren_om.append("Sedimentador 1°")
    elif tipo_primario == "Laguna anaerobia":
        unidades_tren_om.append("Laguna Anaerobia")
    else:
        unidades_tren_om.append("Reactor UASB")
    unidades_tren_om.extend(["Laguna Facultativa", "Lagunas de Maduración"])

    cronograma = []
    for u in unidades_tren_om:
        if u in OM_DATA:
            for p in OM_DATA[u]:
                cronograma.append({"Unidad": u, "Parámetro": p[0],
                                     "Frecuencia": p[1]})
    df_cronograma = pd.DataFrame(cronograma)
    df_freq = df_cronograma.groupby(["Unidad", "Frecuencia"]).size().reset_index(
        name="N° parámetros")
    st.dataframe(df_freq, hide_index=True, use_container_width=True)

    st.markdown("""
    **Recursos humanos sugeridos** (escala Barbosa ~78.000 hab):

    1 **operador líder** con formación en saneamiento (medio tiempo).
    1-2 **auxiliares de operación** para mantenimiento y muestreo (tiempo completo).
    Servicios de **laboratorio externo** mensuales para análisis fisicoquímicos
    completos (DBO, DQO, SST, NTK, P, coliformes).
    1 **ingeniero supervisor** trimestral para evaluación de desempeño y ajustes.
    """)


# -----------------------------------------------------------------
# TAB SENSIBILIDAD
# -----------------------------------------------------------------
with tab_sens:
    st.header("📈 Análisis de sensibilidad")
    st.subheader("🌡️ Sensibilidad a la temperatura")
    T_range = np.arange(15.0, 30.5, 0.5)
    df_T = pd.DataFrame({
        "T (°C)": T_range,
        "Lv anaerobia": [carga_volumetrica_anaerobia(t) for t in T_range],
        "Ls facultativa": [carga_superficial_facultativa(t) for t in T_range],
        "Ef DBO anaerobia (%)": [eficiencia_anaerobia_dbo(t) for t in T_range],
        "TRH UASB (h)": [trh_uasb_por_temperatura(t) for t in T_range],
        "Kb coli (d⁻¹)": [k_corregido(2.6, t, 1.19) for t in T_range],
    })
    cs1, cs2 = st.columns(2)
    cs1.line_chart(df_T.set_index("T (°C)")[["Lv anaerobia", "Ls facultativa"]])
    cs2.line_chart(df_T.set_index("T (°C)")[["Ef DBO anaerobia (%)", "TRH UASB (h)"]])
    st.line_chart(df_T.set_index("T (°C)")[["Kb coli (d⁻¹)"]])


# -----------------------------------------------------------------
# TAB RESUMEN
# -----------------------------------------------------------------
with tab_res:
    st.header("📊 Resumen Ejecutivo")
    st.markdown(f"### Tren: **{tipo_primario}** → Facultativa → Maduración")

    unidades_tren = [
        {"nombre": tipo_primario, "V": primary["V"],
          "TRH": (f"{primary['TRH_h']:.1f} h" if "TRH_h" in primary
                  else f"{primary['TRH_d']:.1f} d"),
          "DBO": primary["DBO_out"], "Coli": primary["Coli_out"]},
        {"nombre": "Facultativa", "V": V_fac, "TRH": f"{TRH_fac:.1f} d",
          "DBO": DBO_out_fac, "Coli": Coli_out_fac},
        {"nombre": f"Maduración ×{int(n_lag)}", "V": V_mad_total,
          "TRH": f"{TRH_total:.1f} d",
          "DBO": DBO_out_mad, "Coli": Coli_out},
    ]
    fig_tren = dibujar_tren(unidades_tren)
    st.pyplot(fig_tren)
    plt.close(fig_tren)

    df_caida = pd.DataFrame({
        "Etapa": ["Afluente", f"Efl. {tipo_primario}",
                    "Efl. Facultativa", "Efl. Maduración"],
        "DBO₅": [DBO_in, primary["DBO_out"], DBO_out_fac, DBO_out_mad],
        "DQO": [DQO_in, primary["DQO_out"], DQO_out_fac, DQO_out_mad],
        "SST": [SST_in, primary["SST_out"], SST_out_fac, SST_out_mad],
        "Coliformes": [Coli_in, primary["Coli_out"], Coli_out_fac, Coli_out],
    })
    st.dataframe(df_caida.style.format({
        "DBO₅": "{:.1f}", "DQO": "{:.1f}", "SST": "{:.1f}",
        "Coliformes": "{:.2e}",
    }), hide_index=True, use_container_width=True)

    st.subheader("Eficiencias globales")
    g1, g2, g3, g4 = st.columns(4)
    g1.metric("DBO₅", f"{(1 - DBO_out_mad/DBO_in)*100:.1f}%")
    g2.metric("DQO", f"{(1 - DQO_out_mad/DQO_in)*100:.1f}%")
    g3.metric("SST", f"{(1 - SST_out_mad/SST_in)*100:.1f}%")
    g4.metric("Coli (log)",
                f"{math.log10(Coli_in/Coli_out) if Coli_out > 0 else float('inf'):.1f}")

    st.subheader("📋 Cumplimiento Resolución 0631 de 2015")
    cumpl = verificar_res0631(DBO_out_mad, DQO_out_mad, SST_out_mad, Coli_out)
    df_cumpl = pd.DataFrame(cumpl)
    st.dataframe(df_cumpl.style.format({"Valor": "{:.2e}", "Límite": "{:.2e}"}),
                  hide_index=True, use_container_width=True)

    # Exportación a Word
    st.subheader("📥 Exportar memoria a Word")
    if DOCX_OK:
        if st.button("📄 Generar memoria.docx"):
            with st.spinner("Generando..."):
                contexto = {
                    "municipio": DATOS_BARBOSA["municipio"],
                    "primario": tipo_primario,
                    "datos": {
                        "Población diseño": f"{pob_diseno:,.0f}",
                        "Qmd": f"{Qmd:,.1f} m³/d",
                        "QMH": f"{QMH:,.1f} m³/d",
                        "DBO₅": f"{DBO_in:.0f} mg/L",
                        "DQO": f"{DQO_in:.0f} mg/L",
                        "T agua": f"{T_agua:.1f} °C",
                        "Régimen": regimen,
                        "Área disp.": f"{area_disp:,.0f} m²",
                    },
                    "fig_tren": fig_to_buf(dibujar_tren(unidades_tren)),
                    "memorias": {},
                    "resumen_df": df_caida.applymap(
                        lambda x: f"{x:.2f}" if isinstance(x, float) else x),
                    "cumplimiento": cumpl,
                }
                # Memoria primario
                if tipo_primario == "Sedimentador 1°":
                    fig_p = dibujar_sedimentador(sed["L"], sed["W"], sed["H"],
                                                    sed["n_paralelo"])
                    contexto["memorias"]["Sedimentador 1°"] = {
                        "fig": fig_to_buf(fig_p),
                        "texto": (f"V = {sed['V']:,.0f} m³ · "
                                    f"A = {sed['A']:,.0f} m²\n"
                                    f"L × W = {sed['L']:.1f} × {sed['W']:.1f} m\n"
                                    f"TRH = {sed['TRH_h']:.2f} h"),
                    }
                elif tipo_primario == "Laguna anaerobia":
                    fig_p = dibujar_laguna(ana["L"], ana["W"], ana["H"],
                                              talud=ana["talud"], tipo="Anaerobia")
                    contexto["memorias"]["Laguna Anaerobia"] = {
                        "fig": fig_to_buf(fig_p),
                        "texto": (f"V = {ana['V']:,.0f} m³ · "
                                    f"TRH = {ana['TRH_d']:.2f} d\n"
                                    f"L × W = {ana['L']:.1f} × {ana['W']:.1f} m"),
                    }
                else:
                    fig_p = dibujar_uasb(uasb["forma"], uasb["L"], uasb["W"],
                                            uasb["H"], uasb["D"])
                    contexto["memorias"]["Reactor UASB"] = {
                        "fig": fig_to_buf(fig_p),
                        "texto": (f"V = {uasb['V']:,.0f} m³ · "
                                    f"TRH = {uasb['TRH_h']:.1f} h\n"
                                    f"Q CH₄ = {uasb['Q_CH4']:,.1f} m³/d"),
                    }
                fig_f = dibujar_laguna(L_f, W_f, H_fac, talud=talud_fac,
                                          tipo="Facultativa")
                contexto["memorias"]["Laguna Facultativa"] = {
                    "fig": fig_to_buf(fig_f),
                    "texto": (f"V = {V_fac:,.0f} m³ · TRH = {TRH_fac:.2f} d\n"
                                f"DBO efluente = {DBO_out_fac:.1f} mg/L"),
                }
                fig_m = dibujar_lagunas_serie(L_m, W_m, H_mad, n_lag)
                contexto["memorias"]["Lagunas de Maduración"] = {
                    "fig": fig_to_buf(fig_m),
                    "texto": (f"{int(n_lag)} en serie · TRH/u = {TRH_unit} d\n"
                                f"Coli salida = {Coli_out:.2e} NMP/100mL"),
                }
                # Layout
                primario_layout = {
                    "L": primary["L"], "W": primary["W"],
                    "n_paralelo": primary.get("n_paralelo", 1),
                    "nombre": tipo_primario,
                }
                fig_lay = dibujar_layout_planta(area_disp, primario_layout,
                                                   {"L": L_f, "W": W_f, "n_paralelo": n_fac},
                                                   {"L": L_m, "W": W_m},
                                                   n_lag, A_lecho_total)
                contexto["memorias"]["Layout General"] = {
                    "fig": fig_to_buf(fig_lay),
                    "texto": f"Lote: {area_disp:,.0f} m² · "
                              f"Tren: {A_total:,.0f} m²",
                }
                doc_bytes = generar_word_memoria(contexto)
                if doc_bytes:
                    st.download_button(
                        "⬇️ Descargar memoria.docx", data=doc_bytes,
                        file_name="memoria_PTAR_Barbosa.docx",
                        mime=("application/vnd.openxmlformats-officedocument."
                              "wordprocessingml.document"),
                    )
                    st.success("✓ Memoria generada con esquemas embebidos.")
    else:
        st.info("Para exportar a Word: `pip install python-docx`")


# Footer
st.markdown("---")
st.caption("Desarrollado para *Sistemas de Tratamiento de Aguas Residuales* (UdeA). " \
"Por Juan Pablo Amud Vásquez - Abril 2026"
           "Lógica en `diseno.py` · Tests en `test_diseno.py` · "
           "Mara (1997), von Sperling & Chernicharo (2005), "
           "van Haandel & Lettinga (1994), RAS 2017, Res 0631/2015.")
